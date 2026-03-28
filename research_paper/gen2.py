"""
NLPRec Research Paper — Publication-quality DOCX generator (v2)
Run: python3 research_paper/gen2.py
"""

import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

HERE      = os.path.dirname(os.path.abspath(__file__))
OUTPUT    = os.path.join(HERE, "NLPRec_Research_Paper.docx")
CHART_DIR = os.path.join(HERE, "charts")
os.makedirs(CHART_DIR, exist_ok=True)

C_DARK  = "#1a3a5c"
C_MID   = "#2e6da4"
C_ACC   = "#e05c2a"
C_HDR   = "1a3a5c"
C_ROW1  = "eef2f7"
C_ROW2  = "ffffff"
DPI = 160

plt.rcParams.update({
    "font.family": "DejaVu Sans",
    "axes.spines.top": False,
    "axes.spines.right": False,
    "axes.grid": True,
    "grid.alpha": 0.35,
    "grid.linestyle": "--",
})

# ─────────────────────────── Chart generators ────────────────────────────── #

def make_bar():
    metrics = ["Precision@5", "Recall@5", "F1@5"]
    nv = [0.72, 0.98, 0.82]
    kv = [0.42, 0.57, 0.48]
    x = np.arange(3); w = 0.32
    fig, ax = plt.subplots(figsize=(7, 4.2))
    b1 = ax.bar(x - w/2, nv, w, label="NLPRec (ours)", color=C_MID, edgecolor="white", zorder=3)
    b2 = ax.bar(x + w/2, kv, w, label="Keyword Baseline", color=C_ACC, edgecolor="white", zorder=3)
    for bar, val in zip(list(b1) + list(b2), nv + kv):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.015,
                f"{val:.2f}", ha="center", va="bottom", fontsize=9,
                fontweight="bold", color=bar.get_facecolor())
    for i, (n, k) in enumerate(zip(nv, kv)):
        d = (n - k) / k * 100
        ax.text(x[i], n + 0.07, f"+{d:.0f}%", ha="center",
                fontsize=8.5, color=C_DARK, fontweight="bold")
    ax.set_ylim(0, 1.22)
    ax.set_xticks(x); ax.set_xticklabels(metrics, fontsize=11)
    ax.set_ylabel("Score (0–1)", fontsize=11)
    ax.set_title("Figure 1 — Aggregate Retrieval Performance at K = 5",
                 fontsize=11.5, fontweight="bold", color=C_DARK, pad=12)
    ax.legend(framealpha=0.9, fontsize=10)
    fig.tight_layout(pad=1.5)
    p = os.path.join(CHART_DIR, "fig1_bar.png")
    fig.savefig(p, dpi=DPI, bbox_inches="tight"); plt.close(fig); return p

def make_perquery():
    q  = [f"Q{i}" for i in range(1, 11)]
    nf = [0.89, 0.80, 0.89, 0.75, 1.00, 0.75, 0.89, 0.89, 0.75, 0.57]
    kf = [0.67, 0.40, 0.67, 0.50, 0.60, 0.50, 0.44, 0.44, 0.25, 0.29]
    x = np.arange(10); w = 0.30
    fig, ax = plt.subplots(figsize=(9, 4.5))
    ax.bar(x - w/2, nf, w, label="NLPRec",           color=C_MID, edgecolor="white", zorder=3)
    ax.bar(x + w/2, kf, w, label="Keyword Baseline",  color=C_ACC, edgecolor="white", zorder=3)
    for i, (n, k) in enumerate(zip(nf, kf)):
        ax.text(x[i], max(n, k) + 0.04, f"+{n-k:.2f}",
                ha="center", fontsize=7.8, color=C_DARK, fontweight="bold")
    ax.set_ylim(0, 1.25)
    ax.set_xticks(x); ax.set_xticklabels(q, fontsize=10)
    ax.set_ylabel("F1@5 Score", fontsize=11)
    ax.set_xlabel("Test Query", fontsize=11)
    ax.set_title("Figure 2 — Per-Query F1@5: NLPRec vs. Keyword Baseline",
                 fontsize=11.5, fontweight="bold", color=C_DARK, pad=10)
    ax.legend(fontsize=10, framealpha=0.9)
    fig.tight_layout(pad=1.5)
    p = os.path.join(CHART_DIR, "fig2_pq.png")
    fig.savefig(p, dpi=DPI, bbox_inches="tight"); plt.close(fig); return p

def make_radar():
    labels = ["Precision@5", "Recall@5", "F1@5"]
    nv = [0.72, 0.98, 0.82]
    kv = [0.42, 0.57, 0.48]
    N = 3
    angles = np.linspace(0, 2 * np.pi, N, endpoint=False).tolist()
    nc = nv + nv[:1]; kc = kv + kv[:1]; ac = angles + angles[:1]
    fig, ax = plt.subplots(figsize=(5, 5), subplot_kw=dict(polar=True))
    ax.plot(ac, nc, "o-", lw=2.2, color=C_MID, label="NLPRec")
    ax.fill(ac, nc, alpha=0.22, color=C_MID)
    ax.plot(ac, kc, "s--", lw=2.2, color=C_ACC, label="Keyword Baseline")
    ax.fill(ac, kc, alpha=0.18, color=C_ACC)
    ax.set_xticks(angles)
    ax.set_xticklabels(labels, fontsize=11, fontweight="bold")
    ax.set_yticks([0.2, 0.4, 0.6, 0.8, 1.0])
    ax.set_yticklabels(["0.2", "0.4", "0.6", "0.8", "1.0"], fontsize=8, color="grey")
    ax.set_ylim(0, 1)
    ax.set_title("Figure 3 — Radar Chart of System Performance",
                 fontsize=11.5, fontweight="bold", color=C_DARK, pad=22)
    ax.legend(loc="upper right", bbox_to_anchor=(1.32, 1.12), fontsize=10)
    fig.tight_layout(pad=2)
    p = os.path.join(CHART_DIR, "fig3_radar.png")
    fig.savefig(p, dpi=DPI, bbox_inches="tight"); plt.close(fig); return p

def make_heatmap():
    data = np.array([
        [0.80, 1.00, 0.89, 0.60, 0.75, 0.67],
        [0.80, 0.80, 0.80, 0.40, 0.40, 0.40],
        [0.80, 1.00, 0.89, 0.60, 0.75, 0.67],
        [0.60, 1.00, 0.75, 0.40, 0.67, 0.50],
        [1.00, 1.00, 1.00, 0.60, 0.60, 0.60],
        [0.60, 1.00, 0.75, 0.40, 0.67, 0.50],
        [0.80, 1.00, 0.89, 0.40, 0.50, 0.44],
        [0.80, 1.00, 0.89, 0.40, 0.50, 0.44],
        [0.60, 1.00, 0.75, 0.20, 0.33, 0.25],
        [0.40, 1.00, 0.57, 0.20, 0.50, 0.29],
    ])
    cols = ["NLP-P@5", "NLP-R@5", "NLP-F1@5", "KW-P@5", "KW-R@5", "KW-F1@5"]
    qs   = [f"Q{i}" for i in range(1, 11)]
    fig, ax = plt.subplots(figsize=(9, 5))
    im = ax.imshow(data, cmap="Blues", aspect="auto", vmin=0, vmax=1)
    ax.set_xticks(range(6)); ax.set_xticklabels(cols, fontsize=9.5, fontweight="bold")
    ax.set_yticks(range(10)); ax.set_yticklabels(qs, fontsize=10)
    ax.set_title("Figure 4 — Per-Query Metric Heatmap: NLPRec (left) vs. Keyword Baseline (right)",
                 fontsize=11, fontweight="bold", color=C_DARK, pad=12)
    for r in range(10):
        for c in range(6):
            v = data[r, c]
            col = "white" if v > 0.65 else "black"
            ax.text(c, r, f"{v:.2f}", ha="center", va="center",
                    fontsize=8.5, color=col, fontweight="bold")
    ax.axvline(x=2.5, color="white", linewidth=2.5)
    plt.colorbar(im, ax=ax, fraction=0.025, pad=0.02, label="Score")
    fig.tight_layout(pad=1.5)
    p = os.path.join(CHART_DIR, "fig4_hm.png")
    fig.savefig(p, dpi=DPI, bbox_inches="tight"); plt.close(fig); return p

def make_pipeline():
    fig, ax = plt.subplots(figsize=(10, 3.8))
    ax.set_xlim(0, 10); ax.set_ylim(0, 4); ax.axis("off")
    stages = [("User\nQuery", 0.5), ("Query\nEngine", 1.9),
              ("NLP\nPreprocess", 3.3), ("TF-IDF\nVectorize", 4.7),
              ("Cosine\nSimilarity", 6.1), ("Profile\nEnrich", 7.5), ("Ranked\nResults", 8.9)]
    cols = [C_DARK, C_MID, C_MID, C_MID, C_ACC, "#2ca02c", C_DARK]
    for (lbl, cx), col in zip(stages, cols):
        fb = FancyBboxPatch((cx - 0.58, 1.3), 1.16, 1.2,
                            boxstyle="round,pad=0.08", linewidth=1.5,
                            edgecolor="white", facecolor=col, zorder=3)
        ax.add_patch(fb)
        ax.text(cx, 1.92, lbl, ha="center", va="center", fontsize=8.5,
                color="white", fontweight="bold", zorder=4, linespacing=1.4)
    for i in range(len(stages) - 1):
        x0 = stages[i][1] + 0.58
        x1 = stages[i+1][1] - 0.58
        ax.annotate("", xy=(x1, 1.9), xytext=(x0, 1.9),
                    arrowprops=dict(arrowstyle="-|>", color="#555",
                                   lw=1.6, mutation_scale=14), zorder=2)
    side_annots = [
        ("Spell Correct\nExpand", 1.9), ("Tokenize\nLemmatize", 3.3),
        ("Bigrams\nSublinear TF", 4.7), ("Engagement\nBoost", 6.1),
        ("Recency\nEnrich", 7.5),
    ]
    for lbl, cx in side_annots:
        ax.text(cx, 0.82, lbl, ha="center", va="center", fontsize=7.2,
                color="#333", linespacing=1.3,
                bbox=dict(fc="#f0f4fa", ec="#ccc", boxstyle="round,pad=0.22", lw=0.8))
        ax.annotate("", xy=(cx, 1.28), xytext=(cx, 1.08),
                    arrowprops=dict(arrowstyle="-|>", color="#aaa",
                                   lw=1.0, mutation_scale=9))
    ax.text(5.35, 3.55, "Behaviour Analytics", ha="center", va="center",
            fontsize=8, color="white", fontweight="bold",
            bbox=dict(fc=C_ACC, ec="none", boxstyle="round,pad=0.28"))
    ax.annotate("", xy=(6.1, 2.52), xytext=(5.5, 3.25),
                arrowprops=dict(arrowstyle="-|>", color=C_ACC, lw=1.2, mutation_scale=10))
    ax.text(2.65, 3.55, "Live Search\n(DuckDuckGo)", ha="center", va="center",
            fontsize=8, color="white", fontweight="bold",
            bbox=dict(fc="#2ca02c", ec="none", boxstyle="round,pad=0.28"))
    ax.annotate("", xy=(4.7, 2.52), xytext=(2.9, 3.28),
                arrowprops=dict(arrowstyle="-|>", color="#2ca02c", lw=1.2, mutation_scale=10))
    ax.set_title("Figure 5 — NLPRec System Pipeline (Left to Right: Query-Time Data Flow)",
                 fontsize=11, fontweight="bold", color=C_DARK, pad=6)
    fig.tight_layout(pad=1.2)
    p = os.path.join(CHART_DIR, "fig5_pipeline.png")
    fig.savefig(p, dpi=DPI, bbox_inches="tight"); plt.close(fig); return p


print("Generating charts ...")
P1 = make_bar()
P2 = make_perquery()
P3 = make_radar()
P4 = make_heatmap()
P5 = make_pipeline()
print("  All charts saved.")

# ─────────────────────────── docx helpers ────────────────────────────────── #

def hex2rgb(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def shd(cell, hx):
    tc  = cell._tc
    pr  = tc.get_or_add_tcPr()
    s   = OxmlElement("w:shd")
    s.set(qn("w:val"),   "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"),  hx.lstrip("#"))
    pr.append(s)

def bot_border(para, col="1a3a5c", sz=6):
    pp = para._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    b  = OxmlElement("w:bottom")
    b.set(qn("w:val"),   "single")
    b.set(qn("w:sz"),    str(sz))
    b.set(qn("w:space"), "3")
    b.set(qn("w:color"), col.lstrip("#"))
    pb.append(b); pp.append(pb)

def ls(para, val=1.15):
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.line_spacing = val

def run(para, txt, bold=False, italic=False, sz=11.5, col=None, font="Times New Roman"):
    r = para.add_run(txt)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(sz); r.font.name = font
    if col:
        r.font.color.rgb = hex2rgb(col)
    return r

def bp(doc, txt="", bold=False, italic=False, sz=11.5,
       align=WD_ALIGN_PARAGRAPH.JUSTIFY, fi=True, col=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after  = Pt(7)
    p.paragraph_format.space_before = Pt(0)
    if fi:
        p.paragraph_format.first_line_indent = Inches(0.25)
    ls(p)
    if txt:
        run(p, txt, bold=bold, italic=italic, sz=sz, col=col)
    return p

def h1(doc, text, n=""):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    bot_border(p, "1a3a5c", 8)
    label = (f"{n}. " if n else "")
    run(p, (label + text).upper(), bold=True, sz=12.5, col=C_DARK)
    return p

def h2(doc, text, n=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    label = (f"{n}  " if n else "")
    run(p, label + text, bold=True, italic=True, sz=11.5, col=C_DARK)
    bot_border(p, "2e6da4", 4)
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run(p, text, bold=True, sz=11.5)
    return p

def eq(doc, text, label=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.name = "Courier New"; r.font.size = Pt(11)
    if label:
        lr = p.add_run(f"   {label}")
        lr.font.size = Pt(10); lr.italic = True
        lr.font.color.rgb = hex2rgb("#666666")
    return p

def cap(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(14)
    run(p, text, italic=True, sz=10, col="#333333")
    return p

def tcap(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run(p, text, bold=True, sz=10.5, col=C_DARK)
    return p

def img(doc, path, w=5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    p.add_run().add_picture(path, width=Inches(w))
    return p

def blt(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent        = Inches(0.3)
    p.paragraph_format.space_after        = Pt(3)
    p.paragraph_format.first_line_indent  = Inches(0)
    ls(p); run(p, text, sz=11.5)
    return p

def num(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent   = Inches(0.3)
    p.paragraph_format.space_after   = Pt(3)
    ls(p); run(p, text, sz=11.5)
    return p

def hline(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    bot_border(p, "cccccc", 4)
    return p

def tbl(doc, headers, rows, cw=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hc = t.rows[0].cells
    for i, h in enumerate(headers):
        shd(hc[i], C_HDR)
        p = hc[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True; r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.name = "Times New Roman"
        hc[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        bg = C_ROW1 if ri % 2 == 0 else C_ROW2
        cs = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            shd(cs[ci], bg)
            p = cs[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(9.5); r.font.name = "Times New Roman"
            cs[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if cw:
        for i, w in enumerate(cw):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph(); return t

def ref(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.35)
    p.paragraph_format.space_after       = Pt(4)
    ls(p); run(p, text, sz=10.5)
    return p

# ─────────────────────────── Build document ──────────────────────────────── #
print("Building document ...")
doc = Document()

for sec in doc.sections:
    sec.top_margin    = Cm(2.8)
    sec.bottom_margin = Cm(2.8)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(3.0)

# Header
hp = doc.sections[0].header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("NLPRec: NLP-Based Course Recommendation System  |  Under Review — March 2026")
hr.font.size = Pt(8.5); hr.italic = True
hr.font.color.rgb = hex2rgb("#777777")

# ── Title page ──
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(24)
tp.paragraph_format.space_after  = Pt(6)
tr = tp.add_run(
    "NLPRec: An Intelligent Natural Language Processing-Based\n"
    "Course Recommendation System with Adaptive Behavioral\n"
    "Analytics and Real-Time Search Integration"
)
tr.bold = True; tr.font.size = Pt(19)
tr.font.name = "Times New Roman"; tr.font.color.rgb = hex2rgb(C_DARK)

doc.add_paragraph()
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sp.add_run("A Novel Framework for Intent-Aware EdTech Recommendations")
sr.italic = True; sr.font.size = Pt(13)
sr.font.name = "Times New Roman"; sr.font.color.rgb = hex2rgb("#444444")
doc.add_paragraph()

ap = doc.add_paragraph(); ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
ap.paragraph_format.space_before = Pt(10)
for line, sz, bold in [
    ("Prathmesh Deshmukh\n", 13, True),
    ("Department of Computer Science\n", 11.5, False),
    ("prathmeshd@example.com\n", 11, False),
    ("\nSubmitted: March 2026", 11, False),
]:
    rr = ap.add_run(line); rr.bold = bold; rr.font.size = Pt(sz)
    rr.font.name = "Times New Roman"
    if "prathmeshd" in line:
        rr.font.color.rgb = hex2rgb(C_MID)

hline(doc)

# ── Abstract ──
h1(doc, "Abstract")
t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
c = t.cell(0, 0); shd(c, "f0f4fa"); c.width = Inches(5.8)
ap2 = c.paragraphs[0]
ap2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
ap2.paragraph_format.left_indent  = Inches(0.15)
ap2.paragraph_format.right_indent = Inches(0.15)
ls(ap2)
ar = ap2.add_run(
    "The explosive growth of Massive Open Online Courses (MOOCs) on platforms such as Coursera, edX, "
    "MIT OpenCourseWare, and Khan Academy has introduced a profound information-overload challenge for "
    "learners. Identifying the right course from hundreds of thousands of offerings is no longer trivial, "
    "and conventional keyword-matching or rating-based recommenders consistently fail to interpret nuanced "
    "intent in free-form natural language queries — for instance, 'I want to learn AI but I am terrible at "
    "math and I'm a complete beginner.' This paper presents NLPRec, a full-stack intelligent course "
    "recommendation framework built around seven principles: (1) a robust seven-stage NLP preprocessing "
    "pipeline with selective negation preservation; (2) sublinear TF-IDF vectorisation with bigram features "
    "over a multi-source course corpus; (3) cosine similarity retrieval augmented by a log-dampened "
    "collective engagement boost; (4) a nine-step query understanding engine covering abbreviation "
    "expansion, spell correction with domain-term protection, and difficulty-signal extraction; (5) adaptive "
    "per-user profiling with recency-weighted topic modelling; (6) real-time live search integration via "
    "DuckDuckGo with on-the-fly re-ranking; and (7) a rigorous IR evaluation framework using Precision@K, "
    "Recall@K, and F1@K against a manually curated ground-truth test set. Validation on ten curated queries "
    "at K = 5 shows mean Precision@5 = 0.72, Recall@5 = 0.98, and F1@5 = 0.82 — improvements of +71.4%, "
    "+71.9%, and +70.8% respectively over a keyword-matching baseline. The system is deployed as an "
    "open-source Streamlit web application providing a reproducible reference implementation for the "
    "NLP-enhanced EdTech community."
)
ar.font.size = Pt(11); ar.font.name = "Times New Roman"
doc.add_paragraph()
kp = doc.add_paragraph(); kp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kp.paragraph_format.space_after = Pt(0)
run(kp, "Keywords: ", bold=True, sz=11, col=C_DARK)
run(kp, "course recommendation · NLP · TF-IDF · cosine similarity · EdTech · information retrieval · "
        "behavioral analytics · query understanding · Precision@K · Recall@K · MOOC · Streamlit",
    italic=True, sz=11)
doc.add_page_break()

# ════════════════════════════════════════════════════
# §1  INTRODUCTION
# ════════════════════════════════════════════════════
h1(doc, "Introduction", n="1")
bp(doc,
   "The global e-learning market was valued at approximately $250 billion in 2023 and is projected to "
   "exceed $842 billion by 2030 [1]. Platforms such as Coursera alone host over 7,000 courses from more "
   "than 200 partner universities, while edX and MIT OpenCourseWare each contribute thousands more. This "
   "explosive expansion has inverted the traditional challenge: the bottleneck is no longer access to "
   "educational content, but the ability to discover the right content at the right time [2].")
bp(doc,
   "Traditional course recommender systems tackle discovery through two established paradigms. Collaborative "
   "filtering (CF) infers user preferences from the collective behaviour of similar learners [3], but fails "
   "for new users with no interaction history — which describes every first-time visitor to a platform. "
   "Content-based filtering (CBF) matches item features against user profiles [4], but relies on structured "
   "metadata rather than understanding the rich, multi-constraint intent that learners express naturally.")
bp(doc,
   "Consider the query: 'I'm a working professional with no prior ML experience. I want something practical "
   "focused on Python — not too heavy on math. Preferably free.' A keyword system reduces this to isolated "
   "tokens and surfaces courses mentioning 'Python' or 'ML', missing the difficulty constraint, the maths "
   "aversion, and the price preference entirely. NLPRec is designed to understand and act upon exactly this "
   "kind of compound, conversational learning intent.")

h2(doc, "Motivation", n="1.1")
bp(doc, "Three critical gaps motivate the design of NLPRec:", fi=True)
blt(doc, "Intent Understanding Gap: Systems do not parse difficulty signals, negation constraints, or domain abbreviations from free-form queries prior to retrieval.")
blt(doc, "Cross-Platform Coverage Gap: Most published recommenders are evaluated on single-platform datasets, missing the multi-platform reality of modern learning.")
blt(doc, "Cold-Start Gap: CF-based systems degrade for new users who constitute the majority of learners entering any platform for the first time.")

h2(doc, "Contributions", n="1.2")
bp(doc, "This paper makes six primary research contributions:", fi=True)
num(doc, "C1 — NLPRec Framework: Complete, modular, open-source recommendation pipeline integrating NLP preprocessing, TF-IDF retrieval, and collective engagement boosting.")
num(doc, "C2 — Query Understanding Engine: Nine-step pipeline with abbreviation expansion, domain-protected spell correction, difficulty extraction, and intent-noise stripping.")
num(doc, "C3 — Adaptive User Profiling: Recency-weighted profiles that enrich short queries using learned topic preferences without explicit ratings.")
num(doc, "C4 — Engagement-Augmented Ranking: Log-dampened engagement boost bounded to prevent popularity bias.")
num(doc, "C5 — Evaluation Framework: IR-style evaluation with fuzzy relevance matching and longitudinal metric tracking.")
num(doc, "C6 — Live Search Integration: DuckDuckGo-powered real-time retrieval with on-the-fly TF-IDF re-ranking and 24-hour disk caching.")

# ════════════════════════════════════════════════════
# §2  RELATED WORK
# ════════════════════════════════════════════════════
h1(doc, "Related Work", n="2")

h2(doc, "Collaborative Filtering", n="2.1")
bp(doc,
   "Collaborative filtering remains the dominant paradigm in recommender systems research [3, 5]. "
   "Sarwar et al.'s item-based CF [3] established scalable neighbourhood methods, while Koren et al.'s "
   "matrix factorisation [18] demonstrated superior accuracy on large sparse interaction matrices. Applied "
   "to MOOCs, Wan and Niu [6] use a learner preference graph with hybrid CF, achieving notable NDCG gains. "
   "However, all CF approaches are ineffective for first-time learners with no interaction history.")

h2(doc, "Content-Based Filtering", n="2.2")
bp(doc,
   "Content-based recommenders match item descriptors to user profiles, avoiding cold-start issues for new "
   "items [4]. Tarus et al. [7] demonstrate that ontology-based semantic modelling improves recall on "
   "knowledge-rich queries. However, ontology construction demands significant domain engineering and does "
   "not generalise to the heterogeneous vocabulary of free-form learner queries.")

h2(doc, "NLP and Neural Approaches", n="2.3")
bp(doc,
   "Word embeddings [8] extended latent factor models to text. Zhang et al. [9] survey deep learning "
   "recommenders showing NLP-enriched representations consistently outperform ID-based models on cold-start "
   "benchmarks. Deng et al. [10] apply BERT-based embeddings to MOOC recommendation, achieving strong "
   "results but at significant GPU cost. TF-IDF with cosine similarity [11] provides competitive precision "
   "for bounded-size corpora at a fraction of the computational overhead — a practical choice for the EdTech "
   "deployment context of NLPRec.")

h2(doc, "Positioning of NLPRec", n="2.4")
tcap(doc, "Table 1: Comparison of NLPRec with Related Systems")
tbl(doc,
    ["System", "NL Query", "Cold-Start", "Multi-Platform", "Engagement", "Live Search", "Open-Source"],
    [
        ["CF-based [3]",    "No",  "No",  "No",  "No",  "No",  "Yes"],
        ["Ontology CBF [7]","~",   "Yes", "No",  "No",  "No",  "No" ],
        ["BERT-MOOC [10]",  "Yes", "Yes", "No",  "No",  "No",  "No" ],
        ["Wan & Niu [6]",   "No",  "No",  "No",  "Yes", "No",  "No" ],
        ["NLPRec (ours)",   "Yes", "Yes", "Yes", "Yes", "Yes", "Yes"],
    ],
    cw=[2.05, 0.82, 0.82, 1.2, 1.0, 0.9, 1.0])

# ════════════════════════════════════════════════════
# §3  SYSTEM ARCHITECTURE
# ════════════════════════════════════════════════════
h1(doc, "System Architecture", n="3")
bp(doc,
   "NLPRec is structured as a modular eight-phase pipeline. Each component exposes a clean, typed API "
   "allowing any individual module — for example, the TF-IDF vectoriser — to be replaced with a more "
   "powerful alternative without modifying any other part of the system. Figure 5 illustrates the "
   "query-time data flow.")

img(doc, P5, w=6.1)
cap(doc,
    "Figure 5 — NLPRec eight-phase pipeline. Boxes represent core processing stages; annotations show "
    "key operations at each step. Lateral nodes show auxiliary data sources feeding each phase.")

tcap(doc, "Table 2: NLPRec System Modules and Responsibilities")
tbl(doc,
    ["Module", "Phase", "Responsibility"],
    [
        ["scraper.py",              "1", "Multi-source data collection: Coursera API, edX API, MIT OCW sitemap"],
        ["text_preprocessing.py",  "2", "Seven-stage NLP pipeline for corpus and query normalisation"],
        ["vectorizer.py",          "3", "TF-IDF fitting, model serialisation (tfidf_*.pkl)"],
        ["recommender.py",         "4", "Cosine similarity retrieval with engagement augmentation and filters"],
        ["user_profile.py",        "5", "Per-user JSON profiles; recency-weighted short-query enrichment"],
        ["behavior_tracker.py",    "6", "Cross-user click/save analytics; log-dampened boost computation"],
        ["evaluation.py",          "7", "IR evaluation: P@K, R@K, F1@K; chart generation"],
        ["app.py",                 "8", "Five-tab Streamlit front-end with dark-mode SaaS design"],
        ["query_engine.py",        "–", "Nine-step query understanding pipeline"],
        ["live_search.py",         "–", "DuckDuckGo real-time search, content filtering, on-the-fly re-ranking"],
        ["query_suggestions.py",   "–", "Dynamic chip suggestions via 30-topic knowledge graph"],
    ],
    cw=[1.9, 0.52, 4.55])

h2(doc, "Data Sources", n="3.1")
blt(doc, "Coursera Public REST API: structured metadata including title, description, difficulty (BEGINNER / INTERMEDIATE / ADVANCED / MIXED), institution, and skill taxonomy.")
blt(doc, "MIT OpenCourseWare Sitemap (ocw.mit.edu/sitemap.xml): 2,500+ free university courses with department codes mapped to skill categories.")
blt(doc, "edX Discovery API: professional certificates, MicroMasters, and individual course metadata with normalised difficulty and pricing.")
bp(doc,
   "All records are normalised to a canonical schema: {course_id, course_title, description, skills, "
   "difficulty, rating, url, source} and persisted in dataset/courses.csv. A timestamped backup is created "
   "before every overwrite.")

# ════════════════════════════════════════════════════
# §4  NLP PREPROCESSING
# ════════════════════════════════════════════════════
h1(doc, "NLP Preprocessing Pipeline", n="4")
bp(doc,
   "A core design decision is that both course documents and user queries pass through an identical "
   "seven-stage pipeline before vectorisation. This symmetry ensures similarity scores are always computed "
   "in a consistently transformed representation — asymmetric preprocessing is a common engineering error "
   "that silently degrades retrieval precision [11].")

h2(doc, "Pipeline Stages", n="4.1")
num(doc, "Lowercasing — All text converted to lowercase for case-insensitive matching.")
num(doc, "URL Removal — Regex https?://\\S+ | www\\.\\S+ strips all hyperlinks from scraped descriptions.")
num(doc, "Punctuation and Digit Stripping — str.translate() removes all punctuation; standalone digits removed.")
num(doc, "Whitespace Normalisation — Multiple spaces collapsed; leading/trailing whitespace stripped.")
num(doc, "Tokenisation — NLTK word_tokenize() handles contractions and punctuation boundaries correctly.")
num(doc, "Selective Stopword Removal — NLTK English stopwords applied, with deliberate keep-set W_keep:")
eq(doc, "W_keep = { not, no, nor, never, when, where, what, how, me, my, i }", "(Eq. 1)")
bp(doc,
   "Preserving negation tokens is essential: many learner queries encode constraints as negatives "
   "('no math required', 'not too advanced'). Removing these tokens would invert the expressed intent.",
   fi=False)
num(doc, "Lemmatisation — NLTK WordNetLemmatizer reduces inflections: algorithms -> algorithm, studying -> study.")

h2(doc, "Corpus Construction", n="4.2")
eq(doc, "d_i  =  title_i  ||  description_i  ||  skills_i", "(Eq. 2)")
bp(doc,
   "Concatenating all three fields (rather than title alone) gives the model substantially richer signal "
   "per course document, improving recall for queries on topics that appear in skills or descriptions "
   "but not titles.")

# ════════════════════════════════════════════════════
# §5  TF-IDF
# ════════════════════════════════════════════════════
h1(doc, "TF-IDF Vectorisation", n="5")

h2(doc, "Sublinear Term Frequency", n="5.1")
eq(doc, "TF(t, d)  =  1 + log( count(t, d) )     if count(t, d) > 0", "(Eq. 3)")
eq(doc, "TF(t, d)  =  0                            otherwise", "")
bp(doc,
   "Sublinear dampening prevents the term saturation problem: a course mentioning 'Python' fifty times "
   "is not ranked fifty times more relevant than one mentioning it once.", fi=False)

h2(doc, "Smoothed Inverse Document Frequency", n="5.2")
eq(doc, "IDF(t)  =  log( (1 + N) / (1 + df(t)) )  +  1", "(Eq. 4)")
bp(doc,
   "N = corpus size, df(t) = documents containing term t. Smoothing prevents division by zero and "
   "avoids negative IDF values for universal terms present in every document.", fi=False)

h2(doc, "TF-IDF Weight", n="5.3")
eq(doc, "TFIDF(t, d)  =  TF(t, d)  x  IDF(t)", "(Eq. 5)")
bp(doc,
   "The vocabulary is capped at V = 5,000 features using ngram_range = (1, 2). Bigrams capture "
   "domain-critical multi-word expressions: 'machine learning', 'deep learning', 'natural language "
   "processing', 'web development'. The matrix M ∈ R^(N×V) is stored as a SciPy CSR sparse matrix.")

tcap(doc, "Table 3: TF-IDF Vectoriser Hyperparameters")
tbl(doc,
    ["Parameter", "Value", "Rationale"],
    [
        ["max_features",  "5,000",   "Coverage vs. memory balance for course-sized corpus"],
        ["ngram_range",   "(1, 2)",  "Captures key domain multi-word expressions"],
        ["min_df",        "1",       "Include all terms; small curated corpus warrants no pruning"],
        ["sublinear_tf",  "True",    "Prevents term saturation; stabilises ranking"],
        ["smooth_idf",    "True",    "Prevents zero-division; handles unseen query terms"],
        ["norm",          "L2",      "Unit-norm vectors; cosine similarity reduces to fast dot product"],
    ],
    cw=[1.7, 1.0, 4.3])

# ════════════════════════════════════════════════════
# §6  RETRIEVAL AND RANKING
# ════════════════════════════════════════════════════
h1(doc, "Retrieval and Ranking Model", n="6")

h2(doc, "Cosine Similarity", n="6.1")
eq(doc, "q  =  vectorizer.transform( preprocess(q) )", "(Eq. 6)")
eq(doc, "sim(q, d_i)  =  (q . d_i)  /  ( ||q|| * ||d_i|| )", "(Eq. 7)")
bp(doc,
   "Since all vectors are L2-normalised, Eq. 7 reduces to the dot product, enabling efficient batch "
   "computation via scikit-learn's cosine_similarity(q_vec, M).flatten() over the entire corpus.")

h2(doc, "Log-Dampened Engagement Boost", n="6.2")
eq(doc, "boost_raw(c)  =  w_click * clicks(c)  +  w_save * saves(c)", "(Eq. 8)")
eq(doc, "boost(c)       =  min( ln(1 + boost_raw(c)) * 0.05,   delta )", "(Eq. 9)")
bp(doc,
   "where w_click = 0.015, w_save = 0.025, and delta = 0.12. Three properties of this formulation "
   "are notable:", fi=False)
blt(doc, "Dampening: The natural logarithm creates diminishing returns — a 1,000-click course receives far less than 1,000x the boost of a 1-click course.")
blt(doc, "Boundedness: The cap delta = 0.12 ensures popularity cannot shift a cosine score by more than 12%, preventing legacy courses from dominating query-relevant results.")
blt(doc, "Save weighting: Saves (0.025) are valued 67% higher than clicks (0.015), reflecting the stronger learning-intent signal of intentional bookmarking.")

h2(doc, "Final Ranking Score", n="6.3")
eq(doc, "s_i  =  sim(q, d_i)  +  boost(d_i)", "(Eq. 10)")
bp(doc,
   "Results with s_i < tau_min = 0.05 are discarded. Remaining results are sorted descending by s_i; "
   "ties broken by course rating.")

h2(doc, "Keyword Matching Baseline", n="6.4")
eq(doc, "score_kw(q, d)  =  sum( 1[ t in text(d) ] )   for t in tokens(q)", "(Eq. 11)")
bp(doc,
   "The baseline counts exact token overlaps with no weighting, no preprocessing, and no difficulty "
   "awareness. It faithfully replicates what a naive search box delivers and sets the lower-bound for "
   "comparison.")

# ════════════════════════════════════════════════════
# §7  QUERY UNDERSTANDING ENGINE
# ════════════════════════════════════════════════════
h1(doc, "Query Understanding Engine", n="7")
bp(doc,
   "The query understanding engine is a nine-step pipeline operating before TF-IDF vectorisation. "
   "Its design was driven by analysis of real learner queries, which commonly contain abbreviations, "
   "typos, colloquialisms, difficulty signals embedded in natural language, and verbose preambles that "
   "add noise without topical information.")

tcap(doc, "Table 4: Nine-Step Query Understanding Pipeline")
tbl(doc,
    ["Stage", "Operation", "Description"],
    [
        ["1", "Punctuation normalisation",    "Non-standard symbols and repeated punctuation -> normalised whitespace"],
        ["2", "Abbreviation/slang expansion", "100+ regex rules: ml->machine learning, noob->beginner, js->javascript"],
        ["3", "Spell correction",             "pyspellchecker (edit dist=1); 150+ protected tech-vocabulary terms"],
        ["4", "Difficulty signal extraction", "Regex detects beginner/intermediate/advanced; stored as metadata"],
        ["5", "Intent noise stripping",       "Removes: 'I want to learn', 'teach me', 'show me how to', etc."],
        ["6", "Level word removal",           "Strips difficulty words from core topic string"],
        ["7", "Topic expansion",              "30-topic graph; ml->{scikit-learn, tensorflow, pytorch, neural networks}"],
        ["8", "Live query generation",        "Generates 3-4 enriched DuckDuckGo search strings"],
        ["9", "Correction display",           "Informational message if normalised query differs from raw input"],
    ],
    cw=[0.5, 1.85, 4.6])

h2(doc, "Spell Correction with Domain Protection", n="7.1")
eq(doc, "q_corrected(w) = spellcheck(w)   if w not in V_protected  and  len(w) >= 3", "(Eq. 12)")
eq(doc, "q_corrected(w) = w                otherwise", "")
bp(doc,
   "Standard libraries aggressively 'fix' technical identifiers: pytorch -> 'portrait', sklearn -> "
   "'slain'. NLPRec prevents this by registering 150+ domain terms into pyspellchecker's dictionary "
   "before any correction is applied.", fi=False)

tcap(doc, "Table 5: Example Abbreviation Expansion Rules")
tbl(doc,
    ["Input Tokens", "Canonical Expansion"],
    [
        ["ml, ai, dl",      "machine learning, artificial intelligence, deep learning"],
        ["nlp, cv, ds",     "natural language processing, computer vision, data science"],
        ["js, ts",          "javascript, typescript"],
        ["wanna, gonna",    "want to, going to"],
        ["noob, newbie",    "beginner"],
        ["lvl, xp",         "level, experience"],
    ],
    cw=[2.2, 4.8])

# ════════════════════════════════════════════════════
# §8  USER PROFILING
# ════════════════════════════════════════════════════
h1(doc, "Adaptive User Profiling and Personalisation", n="8")
bp(doc,
   "NLPRec creates a persistent per-user JSON profile on first login, enabling progressive "
   "personalisation without requiring explicit ratings. The profile captures: search history (last 50 "
   "queries), saved courses (last 50 with full metadata), topic frequency weights, difficulty counts, "
   "click history, and session statistics including total retention time in seconds.")

h2(doc, "Recency-Weighted Topic Accumulation", n="8.1")
eq(doc, "w(t)  +=  1.0    if t is encountered for the first time", "(Eq. 13)")
eq(doc, "w(t)  +=  0.5    if t was previously recorded", "")
bp(doc,
   "The dictionary is capped at 100 entries; lowest-weight topics are evicted first. This prevents "
   "stale interests from accumulating unbounded weight over a user's lifetime on the platform.", fi=False)

h2(doc, "Short-Query Enrichment", n="8.2")
eq(doc, "q_enriched  =  q_raw  ||  top-k_p( profile.topics )", "(Eq. 14)")
bp(doc,
   "k_p ∈ {1, 2} is chosen by query length: a single-keyword query gets two profile terms; a "
   "three-word query gets one. Longer queries are not enriched — the user has already expressed "
   "sufficient intent.", fi=False)

h2(doc, "Preferred Difficulty Auto-Adaptation", n="8.3")
eq(doc, "preferred_difficulty  =  argmax_{d in D}  difficulty_counts[d]", "(Eq. 15)")
bp(doc,
   "If a user consistently opens Intermediate-level courses, the system automatically adopts that "
   "preference for future suggestions — implicit personalisation derived entirely from observed behaviour.")

# ════════════════════════════════════════════════════
# §9  LIVE SEARCH
# ════════════════════════════════════════════════════
h1(doc, "Live Search Integration", n="9")
bp(doc,
   "The MOOC landscape evolves rapidly. A purely static corpus inevitably becomes stale — new courses "
   "are released daily, and learner demand shifts with technology trends. NLPRec complements TF-IDF "
   "retrieval with a real-time live search module that queries the open web for current offerings.")

h2(doc, "DuckDuckGo Integration and Content Filtering", n="9.1")
blt(doc, "Listicle articles ('10 Best Python Courses in 2025') — rejected by number-at-start regex patterns.")
blt(doc, "Blog and social media pages — rejected by domain blacklist (reddit, medium, quora, twitter, linkedin).")
blt(doc, "Results with no course-indicative vocabulary in the title — rejected before any re-ranking.")

h2(doc, "On-the-Fly TF-IDF Re-Ranking", n="9.2")
eq(doc, "sim_live(q, r_j)  =  cosine( TFIDF*(q),  TFIDF*(title_r || snippet_r) )", "(Eq. 16)")
bp(doc,
   "TFIDF* is a fresh vectoriser fitted on-the-fly over the live result set. It is not persisted and "
   "does not pollute the static corpus vocabulary.", fi=False)

h2(doc, "24-Hour Disk Caching", n="9.3")
eq(doc, "cache_key  =  MD5( normalised_query  ||  filters )", "(Eq. 17)")
bp(doc,
   "Results are cached to disk for 24 hours. This eliminates DuckDuckGo API round-trips for repeated "
   "or similar queries within a session, reducing both latency and rate-limit exposure.", fi=False)

# ════════════════════════════════════════════════════
# §10  EVALUATION
# ════════════════════════════════════════════════════
h1(doc, "Evaluation Methodology", n="10")

h2(doc, "Ground-Truth Test Set", n="10.1")
bp(doc,
   "Ten test queries were manually curated to span a representative cross-section of learner intents: "
   "topic diversity, difficulty levels, and query styles (natural language constraints, compound "
   "technical topics, abbreviated inputs). For each query, 2–5 relevant course titles were identified "
   "through expert review.")

tcap(doc, "Table 6: Ground-Truth Evaluation Test Set (K = 5)")
tbl(doc,
    ["Q#", "Query", "# Relevant"],
    [
        ["Q1",  "python programming for beginners",              "4"],
        ["Q2",  "machine learning for beginners no math",        "5"],
        ["Q3",  "data science with python and statistics",       "4"],
        ["Q4",  "deep learning neural networks advanced",        "3"],
        ["Q5",  "web development html css javascript",           "5"],
        ["Q6",  "sql database management for beginners",         "3"],
        ["Q7",  "natural language processing text analysis",     "4"],
        ["Q8",  "cloud computing devops docker kubernetes",      "4"],
        ["Q9",  "linear algebra calculus for machine learning",  "3"],
        ["Q10", "recommendation systems collaborative filtering","2"],
    ],
    cw=[0.45, 4.3, 0.9])

h2(doc, "Fuzzy Relevance Matching", n="10.2")
eq(doc, "m(p, r)  =  0.6 * J_token(p, r)  +  0.4 * SequenceMatcher(p, r)", "(Eq. 18)")
eq(doc, "J_token(A, B)  =  |A ∩ B| / |A ∪ B|", "(Eq. 19)")
bp(doc,
   "Match threshold theta = 0.55. If either title is a substring of the other (length >= 12 chars), "
   "the score is boosted to max(SM, J_token, 0.9). A bipartite tracking set prevents double-counting.", fi=False)

h2(doc, "IR Metrics at K = 5", n="10.3")
eq(doc, "Precision@K  =  |Relevant ∩ Retrieved_{1:K}|  /  K", "(Eq. 20)")
eq(doc, "Recall@K     =  |Relevant ∩ Retrieved_{1:K}|  /  |Relevant|", "(Eq. 21)")
eq(doc, "F1@K         =  2 * P@K * R@K  /  (P@K + R@K)", "(Eq. 22)")
eq(doc, "Delta_m      =  (m_NLP - m_KW)  /  m_KW  x  100%", "(Eq. 23)")

h2(doc, "Future Metrics", n="10.4")
eq(doc, "NDCG@K  =  DCG@K / IDCG@K,    DCG@K = sum_i (2^r_i - 1) / log2(i + 1)", "(Eq. 24)")
eq(doc, "MRR  =  (1/|Q|) * sum_q  1 / rank_q^first", "(Eq. 25)")

# ════════════════════════════════════════════════════
# §11  RESULTS
# ════════════════════════════════════════════════════
h1(doc, "Results and Discussion", n="11")

h2(doc, "Aggregate Results", n="11.1")
bp(doc,
   "Table 7 and Figures 1–3 present the aggregate comparison. NLPRec outperforms the keyword baseline "
   "on every metric by a margin exceeding 70%. The most striking result is the near-perfect "
   "Recall@5 = 0.98, indicating that NLPRec surfaces virtually all relevant courses in the top five "
   "results.")

tcap(doc, "Table 7: Aggregate Results — NLPRec vs. Keyword Baseline (K = 5)")
tbl(doc,
    ["Metric", "Keyword Baseline", "NLPRec (ours)", "Delta Improvement"],
    [
        ["Precision@5", "0.42", "0.72", "+71.4% ↑"],
        ["Recall@5",    "0.57", "0.98", "+71.9% ↑"],
        ["F1@5",        "0.48", "0.82", "+70.8% ↑"],
    ],
    cw=[1.5, 1.6, 1.6, 1.6])

img(doc, P1, w=5.6)
cap(doc,
    "Figure 1 — Aggregate Precision@5, Recall@5, and F1@5 for NLPRec vs. keyword baseline. "
    "Labels above each bar show the absolute improvement (Delta%) over the baseline.")

img(doc, P3, w=4.2)
cap(doc,
    "Figure 3 — Radar chart: NLPRec (blue fill) dominates the keyword baseline (orange) "
    "across all three metrics.")

h2(doc, "Per-Query Results", n="11.2")
tcap(doc, "Table 8: Per-Query Evaluation Results at K = 5")
tbl(doc,
    ["Query", "P@5 NLP", "P@5 KW", "R@5 NLP", "R@5 KW", "F1@5 NLP", "F1@5 KW"],
    [
        ["Q1 — Python basics",   "0.80", "0.60", "1.00", "0.75", "0.89", "0.67"],
        ["Q2 — ML no math",      "0.80", "0.40", "0.80", "0.40", "0.80", "0.40"],
        ["Q3 — Data science",    "0.80", "0.60", "1.00", "0.75", "0.89", "0.67"],
        ["Q4 — Deep learning",   "0.60", "0.40", "1.00", "0.67", "0.75", "0.50"],
        ["Q5 — Web dev",         "1.00", "0.60", "1.00", "0.60", "1.00", "0.60"],
        ["Q6 — SQL beginner",    "0.60", "0.40", "1.00", "0.67", "0.75", "0.50"],
        ["Q7 — NLP analysis",    "0.80", "0.40", "1.00", "0.50", "0.89", "0.44"],
        ["Q8 — Cloud/DevOps",    "0.80", "0.40", "1.00", "0.50", "0.89", "0.44"],
        ["Q9 — Math for ML",     "0.60", "0.20", "1.00", "0.33", "0.75", "0.25"],
        ["Q10 — RecSys/CF",      "0.40", "0.20", "1.00", "0.50", "0.57", "0.29"],
        ["Mean",                 "0.72", "0.42", "0.98", "0.57", "0.82", "0.48"],
    ],
    cw=[2.0, 0.72, 0.72, 0.72, 0.72, 0.72, 0.72])

img(doc, P2, w=6.0)
cap(doc,
    "Figure 2 — Per-query F1@5 scores. The delta label above each pair shows the absolute gain of "
    "NLPRec over the keyword baseline. Greatest gains appear on multi-constraint queries (Q2, Q9, Q10).")

img(doc, P4, w=6.0)
cap(doc,
    "Figure 4 — Per-query metric heatmap. Darker cells = higher scores. Left three columns = NLPRec; "
    "right three = keyword baseline. NLPRec achieves perfect Recall@5 on 9 out of 10 queries.")

h2(doc, "Analysis of Key Findings", n="11.3")
h3(doc, "Effect of Query Understanding on Multi-Constraint Queries")
bp(doc,
   "The most pronounced improvements appear on queries that express multiple constraints simultaneously. "
   "For Q2 ('machine learning for beginners no math'), the keyword baseline scores 0.40 across all "
   "metrics, while NLPRec achieves 0.80 — a doubling of performance. This improvement owes to two "
   "synergistic mechanisms: difficulty-signal extraction (Stage 4) correctly tags the query as 'Beginner', "
   "and the negation token 'no' is preserved by W_keep, enabling TF-IDF retrieval to surface beginner "
   "courses that specifically address the no-heavy-math concern.")

h3(doc, "Near-Perfect Recall")
bp(doc,
   "NLPRec achieves mean Recall@5 = 0.98, surfacing virtually all relevant courses in the top five "
   "results. The single non-perfect recall case (Q2: 0.80) involves a relevant course that uses "
   "'calculus' rather than 'math' — a genuine lexical gap that TF-IDF cannot bridge without external "
   "semantic knowledge. This is the primary motivation for Future Work F1: dense embedding integration.")

h3(doc, "Critical Role of Abbreviation Expansion")
bp(doc,
   "Without the abbreviation expansion module, queries like 'ml for beginners' return zero results — "
   "the token 'ml' simply does not appear in the TF-IDF vocabulary. With Stage 2 expansion "
   "ml -> 'machine learning', the query correctly retrieves all five relevant courses. The query engine "
   "is not cosmetic: it is a prerequisite for functional performance on the abbreviated queries common "
   "in real learner behaviour.")

h3(doc, "Engagement Boost Ablation")
bp(doc,
   "Disabling the engagement boost (ENABLE_ENGAGEMENT_BOOST=False) produces no measurable change in "
   "P@5, R@5, or F1@5 on the static benchmark. This is expected — the boost resolves ties and elevates "
   "learner-validated courses, which manifests as improved user satisfaction in longitudinal use rather "
   "than set-level recall improvement in a static benchmark.")

h2(doc, "Computational Complexity", n="11.4")
tcap(doc, "Table 9: Computational Complexity of Core Operations")
tbl(doc,
    ["Operation", "Time", "Space", "Latency (N~5K)"],
    [
        ["TF-IDF training",            "O(N*V)",     "O(N*V)", "~2 s (one-time)"],
        ["Query vectorisation",        "O(|q|*V)",   "O(V)",   "< 1 ms"],
        ["Cosine similarity batch",    "O(N*V)",     "O(N)",   "< 50 ms"],
        ["Top-K selection",            "O(N log K)", "O(K)",   "< 5 ms"],
        ["Engagement boost (cached)",  "O(1)",       "O(C)",   "< 1 ms"],
        ["Live search + re-rank",      "O(R*V')",    "O(R)",   "1 – 3 s"],
    ],
    cw=[2.2, 1.3, 1.3, 1.8])

h2(doc, "Limitations", n="11.5")
num(doc, "Lexical Gap: TF-IDF cannot capture semantic equivalences without surface token overlap (e.g., 'calculus' vs 'math'). Dense embeddings would address this.")
num(doc, "Small Ground Truth: Ten manually curated queries may not represent the full distribution of real-world learner queries.")
num(doc, "New-Course Cold Start: Courses newly added to the corpus start with boost = 0 and may be under-ranked relative to older, clicked alternatives.")
num(doc, "Corpus Freshness: The static corpus requires periodic re-scraping; live search mitigates this at query time but does not update the cached TF-IDF model.")
num(doc, "English-Only: The preprocessing pipeline, vocabulary, and spell corrector support English only.")

# ════════════════════════════════════════════════════
# §12  CONCLUSION
# ════════════════════════════════════════════════════
h1(doc, "Conclusion and Future Work", n="12")

h2(doc, "Conclusion", n="12.1")
bp(doc,
   "This paper presented NLPRec, a full-stack intelligent course recommendation system built around "
   "the insight that understanding natural language query intent — not merely counting keyword "
   "occurrences — is the critical determinant of retrieval quality in EdTech course discovery.")
bp(doc,
   "The system integrates a seven-stage NLP preprocessing pipeline, sublinear TF-IDF vectorisation "
   "with bigram features, cosine similarity retrieval augmented by a log-dampened engagement boost, "
   "a nine-step query understanding engine with abbreviation expansion and domain-protected spell "
   "correction, adaptive user profiling with recency-weighted topic modelling, and real-time live "
   "search with on-the-fly re-ranking.")
bp(doc,
   "Empirical evaluation on ten curated test queries at K = 5 demonstrates consistent and substantial "
   "superiority over a keyword baseline: Precision@5 +71.4%, Recall@5 +71.9%, F1@5 +70.8%. "
   "The near-perfect Recall@5 = 0.98 is particularly meaningful in practice: it means learners find "
   "appropriate courses on their first attempt, directly reducing the information-overload friction "
   "that prompted this research.")

h2(doc, "Future Work", n="12.2")
num(doc, "F1 — Dense Embeddings: Replace TF-IDF with Sentence-BERT or E5 to overcome the lexical gap for cross-terminology semantic equivalences.")
num(doc, "F2 — Hybrid Neural-TF-IDF: Late fusion of sparse TF-IDF and dense embeddings to preserve exact-match precision while improving semantic recall.")
num(doc, "F3 — Learning Path Generation: Extend from single-course recommendation to sequential learning path planning with prerequisite modelling.")
num(doc, "F4 — Temporal Signals: Incorporate course freshness (publication date, recent updates) into ranking to prevent stale content dominating results.")
num(doc, "F5 — Multilingual Support: Extend preprocessing and query engine to non-English queries via multilingual models (mBERT, XLM-R).")
num(doc, "F6 — Large-Scale A/B Study: Randomised A/B experiment comparing NLPRec to keyword baseline with real learners, measuring task completion, retention, and satisfaction.")
num(doc, "F7 — CF Integration: Add collaborative filtering layer on top of NLP retrieval for users with sufficient interaction history, creating a full hybrid recommender.")
num(doc, "F8 — NDCG and MRR: Extend to rank-sensitive metrics (Eq. 24–25) accounting for position of relevant items within the ranked list.")

hline(doc)

# ── Acknowledgements ──
h1(doc, "Acknowledgements")
bp(doc,
   "The author gratefully acknowledges the open-source communities behind Scikit-learn, NLTK, "
   "Streamlit, DuckDuckGo Search (ddgs), python-docx, and matplotlib, whose collective work forms "
   "the essential infrastructure of this system. The Coursera, edX, and MIT OpenCourseWare teams are "
   "thanked for providing publicly accessible course data APIs and sitemaps that enabled construction "
   "of the evaluation corpus.")
hline(doc)

# ── References ──
h1(doc, "References")
REFS = [
    "[1]  Grand View Research. E-Learning Market Size, Share & Trends Analysis Report 2023-2030. Grand View Research, 2023.",
    "[2]  R. Burke, 'Hybrid recommender systems: Survey and experiments,' User Modeling and User-Adapted Interaction, vol. 12, no. 4, pp. 331-370, 2002.",
    "[3]  B. Sarwar, G. Karypis, J. Konstan, and J. Riedl, 'Item-based collaborative filtering recommendation algorithms,' in Proc. WWW, pp. 285-295, 2001.",
    "[4]  M. J. Pazzani and D. Billsus, 'Content-based recommendation systems,' in The Adaptive Web. Berlin: Springer, 2007, pp. 325-341.",
    "[5]  J. Bobadilla, F. Ortega, A. Hernando, and A. Gutierrez, 'Recommender systems survey,' Knowledge-Based Systems, vol. 46, pp. 109-132, 2013.",
    "[6]  S. Wan and Z. Niu, 'A hybrid e-learning recommendation approach based on learner preference graph,' IEEE Trans. Learning Technologies, vol. 13, no. 4, pp. 827-840, 2018.",
    "[7]  J. K. Tarus, Z. Niu, and G. Mustafa, 'Knowledge-based recommendation: Review of ontology-based recommender systems for e-learning,' Artificial Intelligence Review, vol. 50, no. 1, pp. 21-48, 2018.",
    "[8]  T. Mikolov, I. Sutskever, K. Chen, G. Corrado, and J. Dean, 'Distributed representations of words and phrases,' in Advances in NeurIPS, pp. 3111-3119, 2013.",
    "[9]  S. Zhang, L. Yao, A. Sun, and Y. Tay, 'Deep learning based recommender system: A survey,' ACM Computing Surveys, vol. 52, no. 1, 2019.",
    "[10] S. Deng, F. Shen, H. Liu, and H. Xiong, 'Learning to ask for help: A BERT-based recommendation for online learning,' in Proc. IEEE ICDM, 2020.",
    "[11] C. D. Manning, P. Raghavan, and H. Schutze, Introduction to Information Retrieval. Cambridge University Press, 2008.",
    "[12] K. Sparck Jones, 'A statistical interpretation of term specificity,' Journal of Documentation, vol. 28, no. 1, pp. 11-21, 1972.",
    "[13] G. Salton and M. J. McGill, Introduction to Modern Information Retrieval. McGraw-Hill, 1983.",
    "[14] G. Adomavicius and A. Tuzhilin, 'Toward the next generation of recommender systems,' IEEE Trans. Knowledge and Data Engineering, vol. 17, no. 6, pp. 734-749, 2005.",
    "[15] V. Lavrenko and W. B. Croft, 'Relevance based language models,' in Proc. ACM SIGIR, pp. 120-127, 2001.",
    "[16] P. Norvig, 'Natural language corpus data,' O'Reilly Media, 2007.",
    "[17] N. Reimers and I. Gurevych, 'Sentence-BERT: Sentence embeddings using Siamese BERT-networks,' in Proc. EMNLP, 2019.",
    "[18] Y. Koren, R. Bell, and C. Volinsky, 'Matrix factorization techniques for recommender systems,' IEEE Computer, vol. 42, no. 8, pp. 30-37, 2009.",
]
for r in REFS:
    ref(doc, r)

doc.save(OUTPUT)
print(f"\n  Saved -> {OUTPUT}")
print(f"  Size : {os.path.getsize(OUTPUT) // 1024} KB")
