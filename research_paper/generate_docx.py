"""
NLPRec Research Paper — Publication-quality DOCX generator
Run: python3 research_paper/generate_docx.py
"""

import os, io, math
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.ticker as mticker
from matplotlib.patches import FancyBboxPatch
import numpy as np

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

HERE   = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(HERE, "NLPRec_Research_Paper.docx")
CHART_DIR = os.path.join(HERE, "charts");  os.makedirs(CHART_DIR, exist_ok=True)

C_DARK_BLUE="#1a3a5c"; C_MID_BLUE="#2e6da4"; C_ACCENT="#e05c2a"
C_TABLE_HDR="1a3a5c";  C_TABLE_R1="eef2f7";  C_TABLE_R2="ffffff"
DPI=16"""
NLPRec Research Paper — Publication-quality DOCX generator
Run: python3 res"axes.Run: python3 research_paper/id":True,"grid.alpha":0.35,"grid.linestyle":"--"})

# ── Charts ───?i─import matplotlib
?atplotlib.use("?mport matplotlib.py??import matplotlib.patches as m?mport matplotlib.ticker as mticker
??rom matplotlib.patches import Fan??mport numpy as np

from docx import Documenri
from docx import,"Rfrom docx.shared import [0from docx.enum.text import WD_ALIGN_PARAGRAPH, anfrom docx.oxml.ns import qn
from docx.oxml import OxmlElement=afrom docx.oxml import OxmlPRfrom docx.enum.table import WD_Tco
HERE   = os.path.dirname(os.path.abspath(__file__))
OUTPUT = osselOUTPUT = os.path.join(HERE, "NLPRec_Research_Paper  CHART_DIR = os.path.join(HERE, "charts");  os.makedirs(Cte
C_DARK_BLUE="#1a3a5c"; C_MID_BLUE="#2e6da4"; C_ACCENT="#e05c2a"
C_TABLE_HDR="1   C_TABLE_HDR="1a3a5c";  C_TABLE_R1="eef2f7";  C_TABLE_R2="fffff=bDPI=16"""
NLPlor())
    for i,(n,k) in enumerate(zip(nv,kv)):
        d=(n-Run: python3 res"axes.Run: python3 research_paper/id":True,r"
# ── Charts ───?i─import matplotlib
?atplotlib.use("?mport matplotlib.py??import t_x?atplotlib.use("?mport matplotlib.py??imporel??rom matplotlib.patches import Fan??mport numpy as np

from docx import Documenri
from docx import,"Rfrom  
from docx import Documenri
from docx import,"Rfrom doc_BLfrom docx import,"Rfrom dfrfrom docx.oxml import OxmlElement=afrom docx.oxml import OxmlPRfrom docx.enum.table import WD_Tco
HERE   = os.path.dirnaDPHERE   = os.path.dirname(os.path.abspath(__file__))
OUTPUT = osselOUTPUT = os.path.join(HERE, "NanOUTPUT = osselOUTPUT = os.path.join(HERE, "NLPRec_89C_DARK_BLUE="#1a3a5c"; C_MID_BLUE="#2e6da4"; C_ACCENT="#e05c2a"
C_TABLE_HDR="1   C_TABLE_HDR="1a3a5c";  C_TABLE_R1="eef2f7"loC_TABLE_HDR="1   C_TAB  ax.bar(x-w/2,nf,w,label="NLPRec",color=CNLPlor())
    for i,(n,k) in enumerate(zip(nv,kv)):
        d=(n-Run: python3 res"axes.Rur=    for ,e        d=(n-Run: python3 res"axes.Run: k)# ── Charts ───?i─import matplotlib
?atplotlib.use("?mpo2f?atplotlib.use("?mport matplotlib.py??imporlo
from docx import Documenri
from docx import,"Rfrom  
from docx import Documenri
from docx import,"Rfrom doc_BLfrom docx import,"Rfrom dfrfrom docx.oxml impx.sfrom docx import,"Rfrom  ntfrom docx import Documene(from docx import,"Rfrom d1@HERE   = os.path.dird Baseline",
                 fontsize=11.5,fontweight="bold",color=C_DARK_BLUE,pad=10)
    ax.legend(fontsize=10,framealpha=0.9)
    fOUTPUT = osselOUTPUT = os.path.join(HERE, "NanOUTPUT = osselOUTPUT = os.pasaC_TABLE_HDR="1   C_TABLE_HDR="1a3a5c";  C_TABLE_R1="eef2f7"loC_TABLE_HDR="1   C_TAB  ax.bar(x-w/2,nf,w,label="NLPRec",color=CNLPlor())
    for i,(n,k) in enumer,0    for i,(n,k) in enumerate(zip(nv,kv)):
        d=(n-Run: python3 res"axes.Rur=    for ,e        d=(n-Run: python3 res"axes.Run: k)ig        d=(n-Run(figsize=(5,5),subplot_kw=?atplotlib.use("?mpo2f?atplotlib.use("?mport matplotlib.py??imporlo
from docx import Documenri
from docx import,"Rfrom  
from docxt(from docx import Documenri
from docx import,"Rfrom  
from docx import Docu,kfrom docx import,"Rfrom  NTfrom docx import Documenesfrom docx import,"Rfrom del                 fontsize=11.5,fontweight="bold",color=C_DARK_BLUE,pad=10)
    ax.legend(fontsize=10,framealpha=0.9)
    fOUTPUT = osselOUTPUT = os.path.join(HERE, "NanOUTPUT = osselOUTPUTtle    ax.legend(fontsize=10,framealpha=0.9)
    fOUTPUT = osselOUTPUT = os.ize=11.5,fontweight="bold",color=C_DARK_BLUE,pad=22)
    ax.legend(loc="upper right",bbox_to_anchor=(1.32,1.12),fontsize=10)
    fig.tight_layout(pad=2); p=os.path.join(CHART_DIR,"fig3_radar.png")
    fig.savefig(p,dpi=DPI,bbox_inches="tight"); plt.close(fig);        d=(n-Run: python3 res"axes.Rur=    for ,e        d=(n-.89,0.6from docx import Documenri
from docx import,"Rfrom  
from docxt(from docx import Documenri
from docx import,"Rfrom  
from docx import Docu,kfrom docx import,"Rfrom  NTfrom docx import Documenesfrom docx import,0from docx import,"Rfrom    from docxt(from docx imp0.from docx import,"Rfrom  
from docx .0from docx import Docu,kf      ax.legend(fontsize=10,framealpha=0.9)
    fOUTPUT = osselOUTPUT = os.path.join(HERE, "NanOUTPUT = osselOUTPUTtle    ax.legend(fontsize=10,framealpha=0.9)
    fOUTPUT = osseNL    fOUTPUT = osselOUTPUT = os.path.joinF1    fOUTPUT = osselOUTPUT = os.ize=11.5,fontweight="bold",color=C_DARK_BLUE,pad=22)
    ax.legend(loc="upper rights"    ax.legend(loc="upper right",bbox_to_anchor=(1.32,1.12),fontsize=10)
    fig.ti,f    fig.tight_layout(pad=2); p=os.path.join(CHART_DIR,"fig3_radar.png"ic    fig.savefig(p,dpi=DPI,bbox_inches="tight"); plt.close(fig);       c from docx import,"Rfrom  
from docxt(from docx import Documenri
from docx import,"Rfrom  
from docx import Docu,kfrom docx import,"Rfrom  NTfrom docx import   from docxt(from docx imp       v=data[r,c]; col="white" if v>0.6from docx import Docu,kf  from docx .0from docx import Docu,kf      ax.legend(fontsize=10,framealpha=0.9)
    fOUTPUT = osselOUTPUT = os.path.join(HERE, "NanOUTPUT = osselOUTPUTtle    ax.legend(fonion=0.025,pad=0.02,label="Score")
    fig.tight_layout(pad=1.5); p=os.path.join(CH    fOUTPUT 4_hm.png")
    fig.savefig(p,dpi=DPI,bbox_inches="tight"); plt.close(fig); return p

def make_pipeline()    ax.legend(loc="upper rights"    ax.legend(loc="upper right",bbox_to_anchor=(1.32,1.12),fontsize=10)
    fig.ti,f    fig.tight_layout(pad=2);",    fig.ti,f    fig.tight_layout(pad=2); p=os.path.join(CHART_DIR,"fig3_radar.png"ic    fig.savefig(p,lefrom docxt(from docx import Documenri
from docx import,"Rfrom  
from docx import Docu,kfrom docx import,"Rfrom  NTfrom docx import   from docxt(from docx imp       v=data[r,c]; colbofrom docx import,"Rfrom  
from docx lefrom docx import Docu,kf      fOUTPUT = osselOUTPUT = os.path.join(HERE, "NanOUTPUT = osselOUTPUTtle    ax.legend(fonion=0.025,pad=0.02,label="Score")
    fig.tight_layout(pad=1.5); p=os.path.join(CH    fOUTPUT 4_hm.png")
    fig.savefig(p,dpi=DPI,bbox_inches="ti4)    fig.tight_layout(pad=1.5); p=os.path.join(CH    fOUTPUT 4_hm.png")
    fig.savefig(p,dpi=DPI,bbox_inches="tight"); plt.,x    fig.savefig(p,dpi=DPI,bbox_inches="tight"); plt.close(fig); returco
def make_pipeline()    ax.legend(loc="upper rights"    ax.legend(loc="ct\    fig.ti,f    fig.tight_layout(pad=2);",    fig.ti,f    fig.tight_layout(pad=2); p=os.path.join(CHART_DIR,"fig3_radar.pgafrom docx import,"Rfrom  
from docx import Docu,kfrom docx import,"Rfrom  NTfrom docx import   from docxt(s[1],s[2],s[0]) for s in side]:
        ax.text(cx,cy,txt,ha="center",va="centfrom docx import Docu,kf#3from docx lefrom docx import Docu,kf      fOUTPUT = osselOUTPUT = os.path.join(HERE, "NanOUTPUT = osselOUTPUTtle    ax.legend(fonion=0.025,padyt    fig.tight_layout(pad=1.5); p=os.path.join(CH    fOUTPUT 4_hm.png")
    fig.savefig(p,dpi=DPI,bbox_inches="ti4)    fig.tight_layout(pad=1.5); p=os.path.join(Cte    fig.savefig(p,dpi=DPI,bbox_inches="ti4)    fig.tight_layout(pad=1",    fig.savefig(p,dpi=DPI,bbox_inches="tight"); plt.,x    fig.savefig(p,dpi=DPI,bbox_inches="tight"); plt.close(
 def make_pipeline()    ax.legend(loc="upper rights"    ax.legend(loc="ct\    fig.ti,f    fig.tight_layout(pad=2);",    fig.th\from docx import Docu,kfrom docx import,"Rfrom  NTfrom docx import   from docxt(s[1],s[2],s[0]) for s in side]:
        ax.text(cx,cy,txt,ha="center",va="centfrom docx import Docu,kf#3from docx lefrom docx import Docu,kf          ax.text(cx,cy,txt,ha="center",va="centfrom docx import Docu,kf#3from docx lefrom docx import Docu,kf  LP    fig.savefig(p,dpi=DPI,bbox_inches="ti4)    fig.tight_layout(pad=1.5); p=os.path.join(Cte    fig.savefig(p,dpi=DPI,bbox_inches="ti4)    fig.tight_layout(pad=1",    fig.savefig(p,dpi=DPI,bbox_inches="tight"); plt.,x    fig.savefig(p,dpi=DPI,bbox_inches="tight"); plt.close(
 def makeri def make_pipeline()    ax.legend(loc="upper rights"    ax.legend(loc="ct\    fig.ti,f    fig.tight_layout(pad=2);",    fig.th\from docx import Docu,kfrom docx import,"Rfrom  NTfrom docx import   from docxt(s[1],s[2],s[0]) for s in side]:
        ax.text(cx,cy,txt,ha="cente??        ax.text(cx,cy,txt,ha="center",va="centfrom docx import Docu,kf#3from docx lefrom docx import Docu,kf          ax.text(cx,cy,txt,ha="center",va="centfrom docx import Docu,kf#3from docx lefrom docx import Docu,kf  LP    fig.savefig   def makeri def make_pipeline()    ax.legend(loc="upper rights"    ax.legend(loc="ct\    fig.ti,f    fig.tight_layout(pad=2);",    fig.th\from docx import Docu,kfrom docx import,"Rfrom  NTfrom docx import   from docxt(s[1],s[2],s[0]) for s in side]:
        ax.text(cx,cy,txt,ha="cente??        ax.text(cx,cy,txt,ha="center",va="centfrom docx import Docu,kf#3from docx lefrom docx import Docu,kf          ax.text(cx,cy,txt,ha="center",va="centfrom docx import Docu,kf#3from docx lefrom docx imporE_        ax.text(cx,cy,txt,ha="cente??        ax.text(cx,cy,txt,ha="center",va="centfrom docx import Docu,kf#3from docx lefrom docx import Docu,kf          ax.text(cx,cy,txt,ha="center",va="centfrom docx import Docu,kf#3from docx lefrom docx impor   i        ax.text(cx,cy,txt,ha="cente??        ax.text(cx,cy,txt,ha="center",va="centfrom docx import Docu,kf#3from docx lefrom docx import Docu,kf          ax.text(cx,cy,txt,ha="center",va="centfrom docx import Docu,kf#3from docx lefrom docx imporE_        ax.text(cx,cy,txt,ha="cente??        ax.text(cx,cy,txt,ha="center",va="centfrom docx import Docu,kf#3from docx lefrom docx import Docu,kf          ax.text(cx,cy,txt,ha="center",va="centfrom docx import Docu,kf#3from docx lefrom docx impor   i        ax.text(cx,cy,txt,lignment=WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(6)
    _bot_border(p,"1a3a5c",8)
    lbl=(f"{n}. " if n else ""); run(p,(lbl+text).upper(),bold=True,sz=12.5,col=C_DARK_BLUE); return p

def h2(doc, text, n=""):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
    lbl=(f"{n}  " if n else ""); run(p,lbl+text,bold=True,italic=True,sz=11.5,col=C_DARK_BLUE)
    _bot_border(p,"2e6da4",4); return p

def h3(doc, text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)
    run(p,text,bold=True,sz=11.5); return p

def eq(doc, text, label=""):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(text); r.font.name="Courier New"; r.font.size=Pt(11)
    if label:
        lr=p.add_run(f"   {label}"); lr.font.size    p.paragraph_forTrue
           _bot_border(p,"1a3a5c",8)
    lbl=(f"{n}. " if n else ""); run(p,(lbl+text)do    lbl=(f"{n}. " if n else nt
def h2(doc, text, n=""):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p.paragrPt(    p=doc.add_paragraphic    p.paragraph_format.space_before=Pt(10); p.paragraph_forma      lbl=(f"{n}  " if n else ""); run(p,lbl+text,bold=True,italic=True,sz=11.5,crm    _bot_border(p,"2e6da4",4); return p

def h3(doc, text):
    p=doc.add_paragraph(); p.para,c
def h3(doc, text):
    p=doc.add_para, p    p=doc.add_parp=    run(p,text,bold=True,sz=11.5); return p

def eq(doc, text, label=""):
    p=doc.ace_before=Pt(8); p
def eq(doc, text, label=""):
    p=doc.adadd    p=doc.add_paragraph(); th    p.paragraph_format.space_before=Pt(4); p.paragraph_format.s_p    r=p.add_run(text); r.font.name="Courier New"; r.font.size=Pt(11)
    if lal*    if label:
        lr=p.add_run(f"   {label}"); lr.font.size    t.        lr=pnd           _bot_border(p,"1a3a5c",8)
    lbl=(f"{n}. " if n else ""); r):    lbl=(f"{n}. " if n else ""); rut def h2(doc, text, n=""):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRApa    p=doc.add_paragraph r    p.paragrPt(    p=doc.add_paragraphic    p.paragraph_formaar
def h3(doc, text):
    p=doc.add_paragraph(); p.para,c
def h3(doc, text):
    p=doc.add_para, p    p=doc.add_parp=    run(p,text,bold=True,sz=11.5); return p

def eq(doc, text, label=""):
    p=doc.ace_before=Pt(8); p
def)
     p=doc.add_par Gdef h3(doc, text):
    p=doc.add_pNT    p=doc.add_parro
def eq(doc, text, label=""):
    p=doc.ace_before=Pt(8); p
def eq(doc, text, labehc[    p=doc.ace_before=Pt(8);ntdef eq(doc, text, label=""):      p=doc.adadd    p=doc.ad=T    if lal*    if label:
        lr=p.add_run(f"   {label}"); lr.font.size    t.        lr=pnd           _bot_border(p,"1a3a5c",8)
    lbl=(f"{n}. " if n else ""); r):    lbl=(er        lr=p.add_run(f"_T    lbl=(f"{n}. " if n else ""); r):    lbl=(f"{n}. " if n else ""); ru  for ci,val in enumerate(row):
       p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRApa    p=doc.add_paragraph r    p.paragrPt  def h3(doc, text):
    p=doc.add_paragraph(); p.para,c
def h3(doc, text):
    p=doc.add_para, p    p=doc.add_parp=    run(p,text,bold=True,sz=11f     p=doc.add_par,wdef h3(doc, text):
    p=doc.add_p r    p=doc.add_par.c
def eq(doc, text, label=""):
    p=doc.ace_before=Pt(8); p
def)
     p=doc.add_pa p=    p=doc.ace_before=Pt(8);ardef)
     p=doc.add_par Gdefch   0.    p=doc.add_pNT    p=doc.add_parro
_idef eq(doc, text, label=""):
    p=_f    p=doc.ace_before=Pt(8);s(def eq(doc, text, labehc[   ur        lr=p.add_run(f"   {label}"); lr.font.size    t.        lr=pnd           _bot_border(p,"1a3a5c",8)
    lbl=(f"{n}. " if n else ""??   lbl=(f"{n}. " if n else ""); r):    lbl=(er        lr=p.add_run(f"_T    lbl=(f"{n}. " if n else "");ct       p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRApa    p=doc.add_paragraph r    p.paragrPt  def h3(doc, text):
    p=doc.add_paragraph(); p.para,c
def h3(doc, text)me    p=doc.add_paragraph(); p.para,c
def h3(doc, text):
    p=doc.add_para, p    p=doc.add_parp=    run(p,text,bold=Tru02def h3(doc, text):
    p=doc.add_pc=    p=doc.add_paror    p=doc.add_p r    p=doc.add_par.c
def eq(doc, text, label=""):
    p=doc.ace_before=PtGRAPH.CENTER
tp.paragdef eq(doc, text, label=""):
    p=.p    p=doc.ace_befpace_after=Pdef)
     p=doc.add_pa p=   n    el     p=doc.add_par Gdefch   0.  ng-Based\nCourse Recom_idef eq(doc, text, label=""):
    p=_f    p=doc.ace_before=Pt(8)e     p=_f    p=doc.ace_before=ru    lbl=(f"{n}. " if n else ""??   lbl=(f"{n}. " if n else ""); r):    lbl=(er        lr=p.add_run(f"_T    lbl=(f"{n}. " if n else "");ct       p=doc.add_paragraph(); p.aEN    p=doc.add_paragraph(); p.para,c
def h3(doc, text)me    p=doc.add_paragraph(); p.para,c
def h3(doc, text):
    p=doc.add_para, p    p=doc.add_parp=    run(p,text,bold=Tru02def h3(doc, text):
    p=doc.add_pc=    p=doc.add_paror    p=doc.add_p r    p=docAGdef h3(doc, text)me    p=doc.add_pacdef h3(doc, text):
    p=doc.add_para, p    p=doc.addmu    p=doc.add_parDe    p=doc.add_pc=    p=doc.add_paror    p=doc.add_p r    p=doc.add_par.c
def eq(doamdef eq(doc, text, label=""):
    p=doc.ace_before=PtGRAPH.CENTER
tp.pardd    p=doc.ace_before=PtGRAPr.tp.paragdef eq(doc, text, label=""me    p=.p    p=doc.ace_befpace_after l     p=doc.add_pa p=   n    el     pD_BLUE    p=_f    p=doc.ace_before=Pt(8)e     p=_f    p=doc.ace_before=ru    lbl=(f"{n}. " if n else ""??   lbl=(f"{n}shdef h3(doc, text)me    p=doc.add_paragraph(); p.para,c
def h3(doc, text):
    p=doc.add_para, p    p=doc.add_parp=    run(p,text,bold=Tru02def h3(doc, text):
    p=doc.add_pc=    p=doc.add_paror    p=doc.add_p r    p=docAGdef h3(doc, text)me    p=doc.add_pacdef h3(doc, text odef h3(doc, text):
    p=doc.add_para, p    p=doc.addIT    p=doc.add_paran    p=doc.add_pc=    p=doc.add_paror    p=doc.add_p r    p=docAGdef h3(doc, text)m      p=doc.add_para, p    p=doc.addmu    p=doc.add_parDe    p=doc.add_pc=    p=doc.add_paror    p=doc.add_p r    p=doc  def eq(doamdef eq(doc, text, label=""):
    p=doc.ace_before=PtGRAPH.CENTER
tp.pardd    p=doc.ace_before=PtGRAPr.tp.paragdef e"i    p=doc.ace_before=PtGRAPH.CENTER
tpertp.pardd    p=doc.ace_before=PtGRAeadef h3(doc, text):
    p=doc.add_para, p    p=doc.add_parp=    run(p,text,bold=Tru02def h3(doc, text):
    p=doc.add_pc=    p=doc.add_paror    p=doc.add_p r    p=docAGdef h3(doc, text)me    p=doc.add_pacdef h3(doc, text odef h3(doc, text):
    p=doc.add_para, p    p=doc.addIT    p=doc.add_paran    p=doc.add_pc=    p=do2)    p=doc.add_par v    p=doc.add_pc=    p=doc.add_paror    p=doc.add_p r    p=docAGdef h3(doc, text)mus    p=doc.add_para, p    p=doc.addIT    p=doc.add_paran    p=doc.add_pc=    p=doc.add_paror    p=doc.add_p r    p=docAGdef h3(doc, textst    p=doc.ace_before=PtGRAPH.CENTER
tp.pardd    p=doc.ace_before=PtGRAPr.tp.paragdef e"i    p=doc.ace_before=PtGRAPH.CENTER
tpertp.pardd    p=doc.ace_before=PtGRAeadef h3(doc, text):
    p=doc.add_para, p    p=doc.add_parp=    run(p,text,bold=Tru02def h3(doc, text):
    p=doc.add_pc=    p=doc.aductp.pardd    p=doc.ace_before=PtGRAd tpertp.pardd    p=doc.ace_before=PtGRAeadef h3(doc, text):
    p=doc.add_para, p    p=nd    p=doc.add_para, p    p=doc.add_parp=    run(p,teet. Val    p=doc.add_pc=    p=doc.add_paror    p=doc.add_p r    p=docAGdef h3(doc, text)m0.    p=doc.add_para, p    p=doc.addIT    p=doc.add_paran    p=doc.add_pc=    p=do2)    p=doc.add_par v    p=doc.add_pc=    p=doc.add_parngtp.pardd    p=doc.ace_before=PtGRAPr.tp.paragdef e"i    p=doc.ace_before=PtGRAPH.CENTER
tpertp.pardd    p=doc.ace_before=PtGRAeadef h3(doc, text):
    p=doc.add_para, p    p=doc.add_parp=    run(p,text,bold=Tru02def h3(doc, text):
    p=doc.add_pc=    p=doc.aductp.pardd    p=doc.ace_before=PtGRAd tpertp.pardd    p=doc.ace_before=PtGRAeadef h3(doc, text):attpertp.pardd    p=doc.ace_before=PtGRAeadef h3(doc, text):
    p=doc.add_para, p    p= r    p=doc.add_para, p    p=doc.add_parp=    run(p,text,boTe    p=doc.add_pc=    p=doc.aductp.pardd    p=doc.ace_before=PtGRAd tpertp.pardd   ng    p=doc.add_para, p    p=nd    p=doc.add_para, p    p=doc.add_parp=    run(p,teet. Val    p=doc.add_pc=    p=doc.add_paroructpertp.pardd    p=doc.ace_before=PtGRAeadef h3(doc, text):
    p=doc.add_para, p    p=doc.add_parp=    run(p,text,bold=Tru02def h3(doc, text):
    p=doc.add_pc=    p=doc.aductp.pardd    p=doc.ace_before=PtGRAd tpertp.pardd    p=doc.ace_before=PtGRAeadef h3(doc, text):attpertp.pardd    p=doc.ace_before=PtGRAeadef h3(doc, text):
    p=doc.add_para, p    p= r    p=doc.add_para, p    p=doc.addtiona    p=doc.add_para, p    p=doc.add_parp=    run(p,text,bo      p=doc.add_pc=    p=doc.aductp.pardd    p=doc.ace_before=PtGRAd tpertp.pardd   ti    p=doc.add_para, p    p= r    p=doc.add_para, p    p=doc.add_parp=    run(p,text,boTe    p=doc.add_pc=    p=doc.aductp.pardd    p=doc.ace_before=PtGRAd tpertp.pardd   ng    p=do colle    p=doc.add_para, p    p=doc.add_parp=    run(p,text,bold=Tru02def h3(doc, text):
    p=doc.add_pc=    p=doc.aductp.pardd    p=doc.ace_before=PtGRAd tpertp.pardd    p=doc.ace_before=PtGRAeadef h3(doc, text):attpertp.pardd    p=doc.ace_before=PtGRAeadef h3(doc, text):
    p=doc.add_para, p    p= r    p=doc.add_para, p    p=doc.addtiona    p=doc.add_par t    p=doc.add_pc=    p=doc.aductp.pardd    p=doc.ace_before=PtGRAd tpertp.pardd   ss    p=doc.add_para, p    p= r    p=doc.add_para, p    p=doc.addtiona    p=doc.add_para, p    p=doc.add_parp=    run(p,text,bo      p=doc.add_pc=    p=doc.aductp.pardd    p=doc.ace_befo      p=doc.add_pc=    p=doc.aductp.pardd    p=doc.ace_before=PtGRAd tpertp.pardd    p=doc.ace_before=PtGRAeadef h3(doc, text):attpertp.pardd    p=doc.ace_before=PtGRAeadef h3(doc, text):
    p=doc.add_para, p    p= r    p=doc.add_para, p    p=doc.addtiona    p=doc.add_par t    p=doc.add_pc=    p=doc.aductp.pardd    p=doc.ace_before=PtGRAd tpertp.pardd   ss    p=doc.add_para, p    p= r    p=doc.add_para, p    p=doc.addtiona    p=doc.add_para, p    p=doc.add_parp=    run(p or dom    p=doc.add_para, p    p= r    p=doc.add_para, p    p=doc.addtiona    p=doc.add_par t    p=doc.add_pc=    p=doc.aductp.pardd    p=doc.ace_before=PtGRAd tpertp.pardd   ss    p=doc.addmu    p=doc.add_para, p    p= r    p=doc.add_para, p    p=doc.addtiona    p=doc.add_par t    p=doc.add_pc=    p=doc.aductp.pardd    p=doc.ace_before=PtGRAd tpertp.pardd   ss    p=doc.add_para, p    p= r    p=doc.add_para, p    p=doc.addtiona    p=doc.add_para, p    p=doc.add_parp=    run(p or dom    p=doc.add_para, p    p= r    p=doc.add_para, p    p=doc.addtiona    p=doc.add_par t    p=doc.add_pc=    p=doc.aductp.pardd    p=doc.ace_before=PtGRAd tpertp.pardd   ss    p=doc.addmu    p=doc.add_para, p    p= r    p=doc.add_para, p    p=doc.reviation expansion, domain-protected spell correction, difficulty extraction, and intent-noise stripping.")
num(doc,"C3 — Adaptive User Profiling: Recency-weighted profiles that enrich short queries using learned topic preferences without explicit ratings.")
num(doc,"C4 — Engagement-Augmented Ranking: Log-dampened engagement boost bounded to prevent popularity bias.")
num(doc,"C5 — Evaluation Framework: IR-style evaluation with fuzzy relevance matching and longitudinal metric tracking.")
num(doc,"C6 — Live Search Integration: DuckDuckGo-powered real-time retrieval with on-the-fly TF-IDF re-ranking and 24-hour disk caching.")

# §2 RELATED WORK
h1(doc,"Related Work",n="2")
h2(doc,"Collaborative Filtering",n="2.1")
bp(doc,"Collaborative filtering remains the dominant paradigm in recommender systems research [3, 5]. Sarwar et al.'s "
       "item-based CF [3] established scalable neighbourhood methods, while Koren et al.'s matrix factorisation [18] "
       "demonstrated superior accuracy on large sparsenum(doc,"C3 — Adaptive User Profiling: Recency-weighted profiles that enrich short queries using learnebrnum(doc,"C4 — Engagement-Augmented Ranking: Log-dampened engagement boost bounded to prevent popularity bias.")
num(doc,"C5 — Evaluation Framework:,"num(doc,"C5 — Evaluation Framework: IR-style evaluation with fuzzy relevance matching and longitudinal metric oinum(doc,"C6 — Live Search Integration: DuckDuckGo-powered real-time retrieval with on-the-fly TF-IDF re-ranking and 24-hpr
# §2 RELATED WORK
h1(doc,"Related Work",n="2")
h2(doc,"Collaborative Filtering",n="2.1")
bp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborative Filterebp(doc,"Collaborative filtering remains d        "item-based CF [3] established scalable neighbourhood methods, while Koren et al.'s matrix factorisation [18] " d       "demonstrated superior accuracy on large sparsenum(doc,"C3 — Adaptive User Profiling: Recency-weighted profiolnum(doc,"C5 — Evaluation Framework:,"num(doc,"C5 — Evaluation Framework: IR-style evaluation with fuzzy relevance matching and longitudinal metric oinum(doc,"C6 — Live Search Integration: DuckDuckGo-powered real-time retrieval with on-the-fly TF-IDF re-ranking and 24ti# §2 RELATED WORK
h1(doc,"Related Work",n="2")
h2(doc,"Collaborative Filtering",n="2.1")
bp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborative Filterebp(doc,"Collaborative filtering remains d        "item-based CF [3] established scalab",h1(doc,"Related Wenh2(doc,"Collaborative Filte[3bp(doc,"Collaborative filtering remains 
 h1(doc,"Related Work",n="2")
h2(doc,"Collaborative Filtering",n="2.1")
bp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborative Filterebp(doc,"Collaborative filtering remains d        "item-based CF [3] established scalab",h1(doc,"Related Wenh2(doc,"Collaborative Filte[3bp(doc,"Collaborative filtering remains 
 h1(doc,"Related Work",n="2")
h2(doc,"Collaborative Filtering",n="2.1")
bp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborative Filterebp(doc,"Collaborative filtering remains d        "item-based CF [3] established scalab",h1(doc,"Related Wenh2(doc,"Collaborative Filte[3bp(doc,"Collaboratih2(doc,"Collaborative Filtew=bp(doc,"Collaborative filtering remains ph h1(doc,"Related Work",n="2")
h2(doc,"Collaborative Filtering",n="2.1")
bp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborative Filterebp(doc,"Collaborative filtering remains d        "item-based CF [3] established scalab",h1(doc,"Relatedilh2(doc,"Collaborative Filter",bp(doc,"Collaborative filtering remains AP h1(doc,"Related Work",n="2")
h2(doc,"Collaborative Filtering",n="2.1")
bp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborative Filterebp(doc,"Collaborative filtering remains d        "item-based CF [3] established scalab",h1(doc,"Relatedemh2(doc,"Collaborative Filter"]bp(doc,"Collaborative filtering remains JSh2(doc,"Collaborative Filtering",n="2.1")
bp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborative Filterebp(doc,"Collaborative filtering remains d        "item-based CF [3] established scalab",h1(doc,"Relatedilh2(doc,"Collaborative Filter",bp(doc,"Collaborative filtering remains AP h1(doc,"Related Work",n="2")
h2(doqubp(doc,"Collaborative filtering remains e_h2(doc,"Collaborative Filtering",n="2.1")
bp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborative Filterebp(doc,"Collaborative filtering remains d        "item-based CF [3] established scalab",h1(doc,"Relatedemh2(doc,"Collaborative Filter"]bp(doc,"Collaborative filtg tibp(doc,"Collaborative filtering remains  Ibp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborative Filterebp(doc,"Collaborative filtering remains d        "item-based CF [3] established scalab",h1(doc,"Relatedilh2(doc,"Collaborative Filter",bp(doc,"Collaborative filtering remains AP h1(doc,"Related Work",n="2")
h2(doqubp(h h2(doqubp(doc,"Collaborative filtering remains e_h2(doc,"Collaborative Filtering",n="2.1")
bp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborative Filterebp(doc,"Collaborative filtering remains d        "item-based CF [3] established scalab",h1(doc,"Relatedemh2(doc,"Cpebp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborats h2(doqubp(h h2(doqubp(doc,"Collaborative filtering remains e_h2(doc,"Collaborative Filtering",n="2.1")
bp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborative Filterebp(doc,"Collaborative filtering remains d        "item-based CF [3] established scalab",h1(doc,"Relatedemh2(doc,"Cpebp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborats h2(doqubp(h h2(doqubp(doc,"Collaborative filtering remains e_h2(doc,"Collaborative Filtering",n="2.1")
bp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborative Filtervebp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborative Filtereblebp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborative Filterebp(doc,"Collaborative filtering remains d        "item-based CF [3] established scalab",h1(doc,"Relatedemh2(doc,"Cpebp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborats h2(doqubp(h h2(doqubp(doc,"Collaborative filtering remains e_h2(doc,"Collaborative Filtering",n="2.1")"bp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborative Filtervebp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborative Filtereblebp(doc,"Collaborative filtering remains the dominse h1(doc,"Related Wheh2(doc,"Collaborative Filterebp(doc,"Collaborative filtering remains d        "item-based CF [3] established scalab",h1(doc,"Relateden="4.2")
eq(doc,"d_i  =  title_i  ‖  description_i  ‖  skills_i","(Eq. 2)")
bp(doc,"Concatenating all three fields (rather than title alone) gives the model substantially richer signal per "
       "course document, improving recall for queries on topics that appear in skills or descriptions but not titles.")

# §5 TF-IDF VECTORISATION
h1(doc,"TF-IDF Vectorisation",n="5")
h2(doc,"Sublinear Term Frequency",n="5.1")
eq(doc,"TF(t, d)  =  1 + log( count(t, d) )     if count(t, d) > 0","(Eq. 3)")
eq(doc,"TF(t, d)  =  0                            otherwise","")
bp(doc,"Sublinear dampening prevents the term saturation problem: a course mentioning 'Python' fifty times is not "
       "ranked fifty times more relevant than one mentioning it once.",fi=False)
h2(doc,"Smoothed Inverse Document Frequency",n="5.2")
eq(doc,"IDF(t)  =  log( (1 + N) / (1 + df(t)) )  +  1","(Eq. 4)")
bp(doc,"N = corpus size, df(t) = number of documents containing term t. Smoothing prevents division by zero "
       "and avoids negative IDF eq(dos bp(doc,"Concatenating all three fields (rather than title alone) givec,       "course document, improving recall for queries on topics that appear in skills or descriptions but not titra
# §5 TF-IDF VECTORISATION
h1(doc,"TF-IDF Vectorisation",n="5")
h2(doc,"Sublinear Term Frequency",n="5.1")
eq(doc,"TF', h1(doc,"TF-nguage processinh2(doc,"Sublinear Term Frequency",ne eq(doc,"TF(t, d)  =  1 + log( count(t, d)iPeq(doc,"TF(t, d)  =  0                            otherwise","")
bp(doc,"Subl
tbp(doc,"Sublinear dampening prevents the term saturation proble"5       "ranked fifty times more relevant than one mentioning it once.",fi=False)
h2(doc,"Smoothed Inverse Documentlth2(doc,"Smoothed Inverse Document Frequency",n="5.2")
eq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + N) / (1 + df(t)) ) ue","Pbp(doc,"N = corpus size, df(t) = number of documents containing ",       "and avoids negative IDF eq(dos bp(doc,"Concatenating all three fields (rather than title alone) giveim# §5 TF-IDF VECTORISATION
h1(doc,"TF-IDF Vectorisation",n="5")
h2(doc,"Sublinear Term Frequency",n="5.1")
eq(doc,"TF', h1(doc,"TF-nguage processinh2(doc,"Sublinear Term Frequency",ne eq(doc,"TF(t, d)  =  1 + log( count(t, d)"h1(doc,"TF-IDF Vectorim(q, h2(doc,"Sublinear Term Frequency",n??eq(doc,"TF', h1(doc,"TF-nguage processinh(dbp(doc,"Subl
tbp(doc,"Sublinear dampening prevents the term saturation proble"5       "ranked fifty times more relevant than one mentioning it once.",fi=False)
h2(doc,"Smoothed Invers ttbp(doc,"Suorh2(doc,"Smoothed Inverse Documentlth2(doc,"Smoothed Inverse Document Frequency",n="5.2")
eq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + ooeq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + N) / (1 + df(t)) ) ue","Pbp(docich1(doc,"TF-IDF Vectorisation",n="5")
h2(doc,"Sublinear Term Frequency",n="5.1")
eq(doc,"TF', h1(doc,"TF-nguage processinh2(doc,"Sublinear Term Frequency",ne eq(doc,"TF(t, d)  =  1 + log( count(t, d)"h1(doc,"TF-IDF Vectorim(q, h2(doc,"Sublinear Term Frequency",n??eq(doc,"TF', h1(doc ?2(doc,"Sublinear Teularity cannot sheq(doc,"TF', h1(doc,"TF-nguage processinhentbp(doc,"Sublinear dampening prevents the term saturation proble"5       "ranked fifty times more relevant than one mentioning it once.",fi=False)
h2(doc,"Smoothed Invers ttbp(doc,"Suorh2(doc,"Smoothed Inverse Documentlth2(doc,"Smoothed nkh2(doc,"Smoothed Invers ttbp(doc,"Suorh2(doc,"Smoothed Inverse Documentlth2(doc,"Smoothed Inverse Document Frequency",n="5.2")
eq(doc,"IDF(t)  = cieq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + ooeq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + N) / (1 + dtch2(doc,"Sublinear Term Frequency",n="5.1")
eq(doc,"TF', h1(doc,"TF-nguage processinh2(doc,"Sublinear Term Frequency",ne eq(doc,"TF(t, d)  =  1 + log( count(t, d)"h1(doc,"TF-IDF Vecto peq(doc,"TF', h1(doc,"TF-nguage processinh. h2(doc,"Smoothed Invers ttbp(doc,"Suorh2(doc,"Smoothed Inverse Documentlth2(doc,"Smoothed nkh2(doc,"Smoothed Invers ttbp(doc,"Suorh2(doc,"Smoothed Inverse Documentlth2(doc,"Smoothed Inverse Document Frequency",n="5.2")
eq(doc,"IDF(t)  = cieq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + ooeq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + N) / (1 + dtch2(doc,"Sublinear Term Frequency",n="5.1")
eq(doc,"TF', h1(doc,"deq(doc,"IDF(t)  = cieq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + ooeq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + N) / (1 + dtch2(doc,"Sublinear Term Frequency",n="5.1")
eq(doc,"TF', h1(doc,"TF-ngu"]eq(doc,"TF', h1(doc,"TF-nguage processinh2(doc,"Sublinear Term Frequency",ne eq(doc,"TF(t, d)  =  1 + log( count(t, d)"h1(doc,"TF-IDF Vecto peq(doc,"TF', h1(doc,"TF-nguage processinh. h2(d leq(doc,"IDF(t)  = cieq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + ooeq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + N) / (1 + dtch2(doc,"Sublinear Term Frequency",n="5.1")
eq(doc,"TF', h1(doc,"deq(doc,"IDF(t)  = cieq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + ooeq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + N) / (1 + dtch2(doc,"Sublinear Term Frequency",n="5.1")cueq(doc,"TF', h1(doc,"deq(doc,"IDF(t)  = cieq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + ooeq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + N) / (1 + dtch2(doc,"Sublinear T qeq(doc,"TF', h1(doc,"TF-ngu"]eq(doc,"TF', h1(doc,"TF-nguage processinh2(doc,"Sublinear Term Frequency",ne eq(doc,"TF(t, d)  =  1 + log( count(t, d)"h1(doc,"TF-IDF Vecto peq(doc,"TF', 5,1.85,4.65])
h2(doc,"Spell Ceq(doc,"TF', h1(doc,"deq(doc,"IDF(t)  = cieq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + ooeq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + N) / (1 + dtch2(doc,"Sublinear Term Frequency",n="5.1")cueq(doc,"TF', h1(doc,"deq(doc,"IDF(t)  = cieq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + ooeq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + N) / (1 + dtch2(doc,"Sublinear T qeq(d, h2(doc,"Spell Ceq(doc,"TF', h1(doc,"deq(doc,"IDF(t)  = cieq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + ooeq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + N) / (1 + dtch2(doc,"Sublinear Term Frequency",n="5.1")cueq(doc,"TF', h1(doc,"deq(doc,"IDF(t)  = cieq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + ooeq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + N) / (1 + dtch2(doc,"Sublinear T qeq(d, h2(doc,"Spell Ceq(doc,"TF', h1(doc,"deq(doc,"IDF(t)  = cieq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + ooeq(doc,"IDF(t)  =  log( (useq(doc,"IDF(t)  =  log( (1 + N)  User Profiling and Personalisation",n="8")
bp(doc,"NLPRec creates a persistent per-user JSON profile on first login, enabling progressive personalisation "
       "without requiring explicit ratings. The profile captures: search history (last 50 queries), saved courses "
       "(last 50 with full metadata), topic frequency weights, difficulty counts, click history, and session "
       "statistics including total retention time in seconds.")
h2(doc,"Recency-Weighted Topic Accumulation",n="8.1")
eq(doc,"w(t)  +=  1.0    if t is encountered for the first time","(Eq. 13)")
eq(doc,"w(t)  +=  0.5    if t was previously recorded","")
bp(doc,"The dictionary is capped at 100 entries; lowest-weight topics are evicted first. This prevents stale "
       "interests from accumulating unbounded weight over a user's lifetime on the platform.",fi=False)
h2(doc,"Short-Query Enrichment",n="8.2")
eq(doc,"q_enriched  =  q_raw  ‖  top-k_p( profile.topics )","(Eq. 14)")
bp(doc,"k_p ∈ {1, 2} is chosen by query length: a singlebp(doc,"NLPRec creates a persistent per- t       "without requiring explicit ratings. The profile captures: search history (last 50 queries), saved coursnt       "(last 50 with full metadata), topic frequency weights, difficulty counts, click history, and session "
   _{       "statistics including total retention time in seconds.")
h2(doc,"Recency-Weighted Topic Accumulation",thh2(doc,"Recency-Weighted Topic Accumulation",n="8.1")
eq(doc,"tueq(doc,"w(t)  +=  1.0    if t is encounteredderived eneq(doc,"w(t)  +=  0.5    if t was previously recorded","")
bp(doc,"The dictvebp(doc,"The dictionary is capped at 100 entries; lowest-wol       "interests from accumulating unbounded weight over a user's lifetime on the platform.",fi=False)
h2(doith2(doc,"Short-Query Enrichment",n="8.2")
eq(doc,"q_enriched  =  q_raw  ‖  top-k_p( profile.topics )" Ieq(doc,"q_enriched  =  q_raw  ‖  top-1"bp(doc,"k_p ∈ {1, 2} is chosen by query length: a singlebp(doc,"NLPRecct   _{       "statistics including total retention time in seconds.")
h2(doc,"Recency-Weighted Topic Accumulation",thh2(doc,"Recency-Weighted Topic Accumulation",n="8.1")
eq(doc,"tueq(doc,"w(t)  +=  1.0    if t is encounteredderived eneq(doc,"w(t)  +=  0.5    if t was previously recorded","")
bp(doc,"The dictvebp(doc,"Theveh2(doc,"Recency-Weighted Topic Accumulation",thh2(doc,"Recency-Weig",eq(doc,"tueq(doc,"w(t)  +=  1.0    if t is encounteredderived eneq(doc,"w(t)  +=  0.5    if t was pngbp(doc,"The dictvebp(doc,"The dictionary is capped at 100 entries; lowest-wol       "interests from accumulating unbounde,fh2(doith2(doc,"Short-Query Enrichment",n="8.2")
eq(doc,"q_enriched  =  q_raw  ‖  top-k_p( profile.topics )" Ieq(doc,"q_enriched  =  q_raw  ‖  top-1"bp(doc,"k_p ∈ {1, 2} is cineq(doc,"q_enriched  =  q_raw  ‖  top-k_p( prr h2(doc,"Recency-Weighted Topic Accumulation",thh2(doc,"Recency-Weighted Topic Accumulation",n="8.1")
eq(doc,"tueq(doc,"w(t)  +=  1.0    if t is encounteredderived eneq(doc,"w(t)  +=  0.5    if t was previously recorded","")
bp(doc,"The dictvebpateeq(doc,"tueq(doc,"w(t)  +=  1.0    if t is encounteredderived eneq(doc,"w(t)  +=  0.5    if t was plebp(doc,"The dictvebp(doc,"Theveh2(doc,"Recency-Weighted Topic Accumulation",thh2(doc,"Recency-Weig",eq(doc,"tueq(doc,"w(tqueq(doc,"q_enriched  =  q_raw  ‖  top-k_p( profile.topics )" Ieq(doc,"q_enriched  =  q_raw  ‖  top-1"bp(doc,"k_p ∈ {1, 2} is cineq(doc,"q_enriched  =  q_raw  ‖  top-k_p( prr h2(doc,"Recency-Weighted Topic Accumulation",thh2(doc,"Recency-Weighted Topic Accumulation",n="8.1")
eq(doc,"tueq(doc,"w(t)  +=  1.0    if t is encounteredderived eneq(doc,"w(t)  +=  0.5   eueq(doc,"tueq(doc,"w(t)  +=  1.0    if t is encounteredderived eneq(doc,"w(t)  +=  0.5    if t was previously recorded","")
bp(doc,"The dictvebpateeq(doc,"tueq(doc,"w(t)  +=  1.0    if t is encounteredderived eneq(doc,"w(t)  +=  0.5    if t was plebp(doc,"The dictvebp(doc,"Theveh2 abp(doc,"The dictvebpateeq(doc,"tueq(doc,"w(t)  +=  1.0    if t is encounteredderived eneq(doc,"w(t)  +=  0.5    if t was .4eq(doc,"tueq(doc,"w(t)  +=  1.0    if t is encounteredderived eneq(doc,"w(t)  +=  0.5   eueq(doc,"tueq(doc,"w(t)  +=  1.0    if t is encounteredderived eneq(doc,"w(t)  +=  0.5    if t was previously recorded","")
bp(doc,"The dictvebpateeq(doc,"tueq(doc,"w(t)  +=  1.0    if t is encounteredderived eneq(doc,"w(t)  +=  0.5    if t was plebp(doc,"The dictvebp(doc,"Theveh2 abp(doc,"The dictvebpateeq(doc,"tueq(doc,"w(t)  +=  1.0    if t is encounteredderived eneq(doc,"w(t)  +=  0.5    if t was .4eq(doc,"tueq(doc,"w(t)  +=  1.0 Eqbp(doc,"The dictvebpateeq(doc,"tueq(doc,"w(t)  +=  1.0    if t is encounteredderived eneq(doc,"w(t)  +=  0.5    if t was plebp(doc,"The dictvebp(doc,"Theveh2 abp(doc,"The dictvebpateeq(doc,"tueq(doc,"w(t)  +=  1×bp(doc,"The dictvebpateeq(doc,"tueq(doc,"w(t)  +=  1.0    if t is encounteredderived eneq(doc,"w(t)  +=  0.5    if t was plebp(doc,"The dictvebp(doc,"Theveh2 abp(doc,"The dictvebpateeq(doc,"tueq(doc,"w(t)  +=  1.0    if t is encounteredderived eneq(doc,"w(t)  +=  0.5    if t was .4eq(doc,"tueq(doc,"w(t)  +=  1.0 Eqbp(doc,"The dictvebpateeq(doc,"tueq(doc,"w(t)  +=  1.0    if t is encounteredderived eneq(doc,"w(t)  +=  0.5    if t was plebp(doc,"The dictvebp(doc,"Theveh2 abp(doc,"The dictves the near-perfect Recall@5 = 0.98, "
       "indicating that NLPRec surfaces virtually all relevant courses in the top five results.")
tcap(doc,"Table 7: Aggregate Results — NLPRec vs. Keyword Baseline (K = 5)")
tbl(doc,["Metric","Keyword Baseline","NLPRec (ours)","Δ Improvement"],
    [["Precision@5","0.42","0.72","+71.4% ↑"],
     ["Recall@5","0.57","0.98","+71.9% ↑"],
     ["F1@5","0.48","0.82","+70.8% ↑"]],
    cw=[1.5,1.6,1.6,1.6])
img(doc,P1,w=5.6)
cap(doc,"Figure 1 — Aggregate Precision@5, Recall@5, and F1@5 for NLPRec vs. keyword baseline. "
        "Labels above each bar show the absolute improvement (Δ%) over the baseline.")
img(doc,P3,w=4.2)
cap(doc,"Figure 3 — Radar chart: NLPRec (blue fill) dominates the keyword baseline (orange) across all three metrics.")
h2(doc,"Per-Query Results",n="11.2")
tcap(doc,"Table 8: Per-Query Evaluation Results at K = 5")
tbl(doc,["Query","P@5 NLP","P@5 KW","R@5 NLP","R@5 KW","F1@5 NLP","F1@5 KW"],
    [["Q1 — Python basics","0.80","0.60","1.00","0.75"       "indicating that NLPRec surnotcap(doc,"Table 7: Aggregate Results — NLPRec vs. Keyword Baseline (K = 5)")
tbl(doc,["Metric"00tbl(doc,["Metric","Keyword Baseline","NLPRec (ours)","Δ Improvement"],
    [.6    [["Precision@5","0.42","0.72","+71.4% ↑"],
     ["Recall@5","0.5"1     ["Recall@5","0.57","0.98","+71.9% ↑"],
 ",     ["F1@5","0.48","0.82","+70.8% ↑"]],
 Q7    cw=[1.5,1.6,1.6,1.6])
img(doc,P1,w=5.0.img(doc,P1,w=5.6)
cap(do [cap(doc,"Figure vO        "Labels above each bar show the absolute improvement (Δ%) over the baseline.20","1.00","0.img(doc,P3,w=4.2)
cap(doc,"Figure 3 — Radar chart: NLPRec (blue fill) dominates the "]cap(doc,"Figure 0.h2(doc,"Per-Query Results",n="11.2")
tcap(doc,"Table 8: Per-Query Evaluation Results at K = 5")
tbl(doc,["Query","P@5 NL2 tcap(doc,"Table 8: Per-Query Evaluaabtbl(doc,["Query","P@5 NLP","P@5 KW","R@5 NLP","R@5 KW","F      [["Q1 — Python basics","0.80","0.60","1.00","0.75"       "indicating tQ2tbl(doc,["Metric"00tbl(doc,["Metric","Keyword Baseline","NLPRec (ours)","Δ Improvement"],
    [.6    [["Precision@5","0.42","0.72","+71.4% ?      "right three = keyword    [.6    [["Precision@5","0.42","0.72","+71.4% ↑"],
     ["Recall@5","0.5"1     ["Rec F     ["Recall@5","0.5"1     ["Recall@5","0.57","0.98",ng ",     ["F1@5","0.48","0.82","+70.8% ↑"]],
 Q7    cw=[1.5,1.6,1.en Q7    cw=[1.5,1.6,1.6,1.6])
img(doc,P1,w=5.onimg(doc,P1,w=5.0.img(doc,P1  cap(do [cap(doc,"Figure vO       orcap(doc,"Figure 3 — Radar chart: NLPRec (blue fill) dominates the "]cap(doc,"Figure 0.h2(doc,"Per-Query Results",n="11.2")
tcap(doc,"Table . tcap(doc,"Table 8: Per-Query Evaluation Results at K = 5")
tbl(doc,["Query","P@5 NL2 tcap(doc,"Table 8: Per-Query Evaluaabtertbl(doc,["Query","P@5 NL2 tcap(doc,"Table 8: Per-Query Ev p    [.6    [["Precision@5","0.42","0.72","+71.4% ?      "right three = keyword    [.6    [["Precision@5","0.42","0.72","+71.4% ↑"],
     ["Recall@5","0.5"1     ["Rec F     ["Recall@5","0.5"1     ["Recall@5","0.57","0.98",ng ",     ["F1@5","0.48","0.82","+70.8% ↑"]],
 Q7    cw=[1.5,1 "     ["Recall@5","0.5"1     ["Rec F     ["Recall@5","0.5"1     ["Recall@5","0.57","0.98",ng ",     ["F1@5","0.48","0.82","+70.8% ↑"]xi Q7    cw=[1.5,1.6,1.en Q7    cw=[1.5,1.6,1.6,1.6])
img(doc,P1,w=5.onimg(doc,P1,w=5.0.img(doc,P1  cap(do [cap(doc,"Figure vO       orcaseimg(doc,P1,w=5.onition.")
h3(doc,"Critical Role of Atcap(doc,"Table . tcap(doc,"Table 8: Per-Query Evaluation Results at K = 5")
tbl(doc,["Query","P@5 NL2 tcap(doc,"Table 8: Per-Query Evaluaabtertbl(doc,["Query","P@5 NL2 tcap(doc,"Table 8: Per-Query Ev p   Wtbl(doc,["Query","P@5 NL2 tcap(doc,"Table 8: Per-Query Evaluaabtertbl(doc,[tl     ["Recall@5","0.5"1     ["Rec F     ["Recall@5","0.5"1     ["Recall@5","0.57","0.98",ng ",     ["F1@5","0.48","0.82","+70.8% ↑"]],
 Q7    cw=[1.5,1 "     ["Recall@5","0.5"1     ["Rec F     ["Recall@5","0.5"1     ["Recall@5","0.57","0.98",ng ",     ["F1@en Q7    cw=[1.5,1 "     ["Recall@5","0.5"1     ["Rec F     ["Recall@5","0.5"1     ["Recall@5","0.57","0.98",ng ",     ["F1@5","0.48","0.xpimg(doc,P1,w=5.onimg(doc,P1,w=5.0.img(doc,P1  cap(do [cap(doc,"Figure vO       orcaseimg(doc,P1,w=5.onition.")
h3(doc,"Critical Role of Atcap(doc,"Table . tcap(doc,"Table 8: Per-Query Evaluation Results aa h3(doc,"Critical Role of Atcap(doc,"Table . tcap(doc,"Table 8: Per-Query Evaluation Results at K = 5")
tbl(dof tbl(doc,["Query","P@5 NL2 tcap(doc,"Table 8: Per-Query Evaluaabtertbl(doc,["Query","P@5 NL2 tcap(doc,(N Q7    cw=[1.5,1 "     ["Recall@5","0.5"1     ["Rec F     ["Recall@5","0.5"1     ["Recall@5","0.57","0.98",ng ",     ["F1@en Q7    cw=[1.5,1 "     ["Recall@5","0.5"1     ["Rec F     ["Recall@5","0.5"1     ["Recall@5","0.57","0.98",ng ",     ["F1@5","0.48","0.xpimg(doc,P1,w=5.onimg(doc,P1,w=5.0.img(doc,P1  cap(do [cap(doc,"Figure vO       o3,h3(doc,"Critical Role of Atcap(doc,"Table . tcap(doc,"Table 8: Per-Query Evaluation Results aa h3(doc,"Critical Role of Atcap(doc,"Table . tcap(doc,"Table 8: Per-Query Evaluation Results at K = 5")
tbl(dof tbl(doc,["Query","P@5 NL2 tcap(doc,"Table 8: Per-Query Evaluaabtertbl(doc,["Query","P@5 NL2 tcap(doc,(N Q7    cw=[1.5,1 "     ["Recall@5","0.5"1     ["Rec F     ["Ry tbl(dof tbl(doc,["Query","P@5 NL2 tcap(doc,"Table 8: Per-Query Evaluaabtertbl(doc,["Query","P@5 NL2 tcap(doc,(N Q7    cw=[1.5,1 "     ["Recall@5","0.5"1     ["Rec F     ["Recall@5","0.5"1     ["Rem(tbl(dof tbl(doc,["Query","P@5 NL2 tcap(doc,"Table 8: Per-Query Evaluaabtertbl(doc,["Query","P@5 NL2 tcap(doc,(N Q7    cw=[1.5,1 "     ["Recall@5","0.5"1     ["Rec F     ["Ry tbl(dof tbl(doc,["Query","P@5 NL2 tcap(doc,"Table 8: Per-Query Evaluaabtertbl(doc,["Query","P@5 NL2 tcap(doc,(N Q7    cw=[1.5,1 "     ["Recall@5","0.5"1     ["Rec F     ["Recall@5","0.5"1     ["Rem(tbl(dof tbl(doc,["Query","P@5 NL2 tcap(doc,"Table 8: Per-Query Evaluaabtertbl(doc,["Query","P@5 NL2 tcap(doc,(N Q7    cw=[1.5,1 "     ["Recall@5","0.5"1     ["Rec F     ["Ry tbl(dof tbl(doc,["Query","P@5 NL2 tcap(doc,"Table 8: Per-Query Evaluaabtertbl(doc,["Query","P@5 NL2 tcaph course discovery.")
bp(doc,"The system integrates a seven-stage NLP preprocessing pipeline, sublinear TF-IDF vectorisation with bigram "
       "features, cosine similarity retrieval augmented by a log-dampened engagement boost, a nine-step query "
       "understanding engine with abbreviation expansion and domain-protected spell correction, adaptive user profiling "
       "with recency-weighted topic modelling, and real-time live search with on-the-fly re-ranking.")
bp(doc,"Empirical evaluation on ten curated test queries at K = 5 demonstrates consistent and substantial superiority "
       "over a keyword baseline: Precision@5 +71.4%, Recall@5 +71.9%, F1@5 +70.8%. The near-perfect Recall@5 = 0.98 "
       "is particularly meaningful in practice: it means learners find appropriate courses on their first attempt, "
       "directly reducing the information-overload friction that prompted this research.")
h2(doc,"Future Work",n="12.2")
num(doc,"F1 — Dense Embeddings: Replace TF-IDF with Sentence-BERT or E5bp(doc,"The systemrc       "features, cosine similarity retrieval augmented by a log-dampened engagement boost, a nine-step query "
   n        "understanding engine with abbreviation expansion and domain-protected spell correction, adaptive user (d       "with recency-weighted topic modelling, and real-time live search with on-the-fly re-ranking.")
bp(doc,"Empiricalprbp(doc,"Empirical evaluation on ten curated test queries at K = 5 demonstrates consistent and substanat       "over a keyword baseline: Precision@5 +71.4%, Recall@5 +71.9%, F1@5 +70.8%. The near-perfect Recall@5 = 0.98 "
xt       "is particularly meaningful in practice: it means learners find appropriate courses on their first attempt, "La       "directly reducing the information-overload friction that prompted this research.")
h2(doc,"Future Work",n="ash2(doc,"Future Work",n="12.2")
num(doc,"F1 — Dense Embeddings: Replace TF-IDF with Sentatnum(doc,"F1 — Dense Embeddi N   n        "understanding engine with abbreviation expansion and domain-protected spell correction, adaptive user (d       "with recency-weighted topic modelling, and real-time live search with on-the-fsibp(doc,"Empiricalprbp(doc,"Empirical evaluation on ten curated test queries at K = 5 demonstrates consistent and substanat       "over a keyword baseline: Precision@5 +71.4%, Recall@5 +71.9%, F1@5 +70.8%. The near-perf  xt       "is particularly meaningful in practice: it means learners find appropriate courses on their first attempt, "La       "directly reducing the information-overload friction that prompted this research.")
h2(doc,"Future Work",n="ashcch2(doc,"Future Work",n="ash2(doc,"Future Work",n="12.2")
num(doc,"F1 — Dense Embeddings: Replace TF-IDF with Sentatnum(doc,"F1 — Dense Embeddi N   n        "understanding engine with abbreviation expansionrenum(doc,"F1 — Dense Embeddings: Replace TF-IDF with Srch2(doc,"Future Work",n="ashcch2(doc,"Future Work",n="ash2(doc,"Future Work",n="12.2")
num(doc,"F1 — Dense Embeddings: Replace TF-IDF with Sentatnum(doc,"F1 — Dense Embeddi N   n        "understanding engine with abbreviation expansionrenum(doc,"F1 — Dense Embeddings: Replace TF-IDF with Srch2(doc,"Future Work",n="ashcch2(doc,"Future Work",n="ash2(doc,"Future Work",n="12.2")
num(doc,"F1 — Dense Embeddings: Replace TF-IDF with Sentatnum(doc,"F1 — Dense Embeddi N   n        "understanding engine with abbreviation expansionrenum(doc,"F1 — Dense Embeddings: Replace TF-IDF with Srch2(doc,"Future Work",n="ashcch2(doc,"Future Work",n="ash2(doc,"Future Work",n="12.2")
num(doc,"F1 — Dense Embeddings: Replace TF- gnum(doc,"F1 — Dense Embeddings: Replace TF-IDF with Sentatnum(doc,"F1 — Dense Em "num(doc,"F1 — Dense Embeddings: Replace TF-IDF with Sentatnum(doc,"F1 — Dense Embeddi N   n        "understanding engine with abbreviation expansionrenum(doc,"F1 — Dense Embeddings: Replace TF-IDF with Srch2(doc,"Future Work",n="ashcch2(doc,"Future Work",n="ash2(doc,"Future Work",n="12.2"ennum(doc,"F1 — Dense Embeddings: Replace TF-IDF with Sentatnum(doc,"F1 — Dense Embeddi N   n        "understanding engine with abbreviation expansionrenum(doc,"F1 — Dense Embeddings: Replace TF-IDF with Srch2(doc,"Future Work",n="ashcch2(doc,"Future Work",n="ash2(doc,"Future Work",n="12.2"ornum(doc,"F1 — Dense Embeddings: Replace TF- gnum(doc,"F1 — Dense Embeddings: Replace TF-IDF with Sentatnum(doc,"F1 — Dense Em "num(doc,"F1 — Dense Embeddings: Replace TF-IDF with Sentatnum(doc,"F1 — Dense Embeddi N   n        "understanding engine with abbrevion of term specificity,' Journal of Documentation, vol. 28, no. 1, pp. 11–21, 1972.",
    "[13] G. Salton and M. J. McGill, Introduction to Modern Information Retrieval. McGraw-Hill, 1983.",
    "[14] G. Adomavicius and A. Tuzhilin, 'Toward the next generation of recommender systems,' IEEE Trans. Knowledge and Data Engineering, vol. 17, no. 6, pp. 734–749, 2005.",
    "[15] V. Lavrenko and W. B. Croft, 'Relevance based language models,' in Proc. ACM SIGIR, pp. 120–127, 2001.",
    "[16] P. Norvig, 'Natural language corpus data,' O'Reilly Media, 2007. [Online]. Available: https://norvig.com/spell-correct.html",
    "[17] N. Reimers and I. Gurevych, 'Sentence-BERT: Sentence embeddings using Siamese BERT-networks,' in Proc. EMNLP, 2019.",
    "[18] Y. Koren, R. Bell, and C. Volinsky, 'Matrix factorization techniques for recommender systems,' IEEE Computer, vol. 42, no. 8, pp. 30–37, 2009.",
]: ref(doc,txt)

doc.save(OUTPUT)
print(f"\n✓  Saved → {OUTPUT}")
print(f"   Size: {os.path.getsize(OUTPUT)//1024} KB")
